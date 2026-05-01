#!/usr/bin/env python3
"""Generate HEALTH AI SRS Document — comprehensive final version with diagrams."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ───────── PAGE SETUP ─────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ───────── STYLES ─────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

SCREENSHOTS_DIR = os.path.join(os.path.dirname(__file__), 'screenshots')
DIAGRAMS_DIR = os.path.join(os.path.dirname(__file__), 'diagrams')


def add_screenshot(filename, caption, width=Inches(5.8)):
    path = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph(f'Figure: {caption}')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()


def add_diagram(filename, caption, width=Inches(6.2)):
    path = os.path.join(DIAGRAMS_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph(f'Diagram: {caption}')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()


def add_table(headers, rows, header_size=10, body_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(header_size)
    for rd in rows:
        r = t.add_row()
        for i, tx in enumerate(rd):
            r.cells[i].text = str(tx)
            r.cells[i].paragraphs[0].runs[0].font.size = Pt(body_size)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HEALTH AI')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x3b, 0x3b, 0x98)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('European HealthTech Co-Creation & Innovation Platform')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

doc_title = doc.add_paragraph()
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = doc_title.add_run('Software Requirements Specification (SRS)')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Version 3.0 — Final Release • April 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

group_title = doc.add_paragraph()
group_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = group_title.add_run('Group: Health Shield')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

members_p = doc.add_paragraph()
members_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for member in ['Cem Özal', 'Emre Kurubaş', 'Hasabu Can Eltayeb', 'Sertaç Ataç']:
    run = members_p.add_run(member + '\n')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

info_table = doc.add_table(rows=0, cols=2)
info_table.style = 'Light Grid Accent 1'
info_data = [
    ('Document', 'Software Requirements Specification'),
    ('Project', 'HEALTH AI Platform'),
    ('Group', 'Health Shield'),
    ('Members', 'Cem Özal, Emre Kurubaş, Hasabu Can Eltayeb, Sertaç Ataç'),
    ('Version', '3.1 (Final)'),
    ('Date', 'May 1, 2026'),
    ('Course', 'SENG 384 — Software Engineering'),
    ('Standard', 'IEEE 830-1998 / ISO/IEC/IEEE 29148:2018'),
    ('Status', 'Final — Comprehensive with UML diagrams'),
]
for label, value in info_data:
    row = info_table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# REVISION HISTORY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('Revision History', level=1)
add_table(
    ['Version', 'Date', 'Author(s)', 'Description of Changes'],
    [
        ('1.0', '2025-12-10', 'Health Shield', 'Initial SRS — base requirements & system context'),
        ('1.5', '2026-02-14', 'Health Shield', 'Added GDPR, audit trail and admin requirements'),
        ('2.0', '2026-04-01', 'Health Shield', 'Updated UI section to match production state, added screenshots'),
        ('3.0', '2026-04-27', 'Health Shield', 'Final — added UML diagrams (Use Case, Context, DFD, ERD, State, Architecture), expanded NFRs, traceability matrix, glossary'),
    ],
    body_size=9,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '   1.1 Purpose',
    '   1.2 Scope',
    '   1.3 Definitions, Acronyms, and Abbreviations',
    '   1.4 References',
    '   1.5 Document Conventions',
    '   1.6 Intended Audience',
    '   1.7 Overview',
    '2. Overall Description',
    '   2.1 Product Perspective',
    '   2.2 System Context Diagram',
    '   2.3 Product Functions',
    '   2.4 User Classes and Characteristics',
    '   2.5 Operating Environment',
    '   2.6 Design and Implementation Constraints',
    '   2.7 Assumptions and Dependencies',
    '3. Use Cases',
    '   3.1 Use Case Diagram',
    '   3.2 Use Case Specifications',
    '4. Specific Requirements',
    '   4.1 Functional Requirements',
    '   4.2 Non-Functional Requirements',
    '   4.3 External Interface Requirements',
    '   4.4 User Interface Requirements',
    '5. System Models',
    '   5.1 Data Flow Diagram (Level 1)',
    '   5.2 Entity-Relationship Diagram',
    '   5.3 State Transition Diagram',
    '6. System Architecture',
    '   6.1 Architectural Overview',
    '   6.2 Technology Stack',
    '7. Data Model',
    '   7.1 Users Collection',
    '   7.2 Posts Collection (with subcollections)',
    '   7.3 Conversations & Messages',
    '   7.4 Activity Logs & Notifications',
    '8. Requirements Traceability Matrix',
    '9. Acceptance Criteria & Verification',
    '10. Appendix A — UI Screenshots',
    '11. Appendix B — Glossary',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph(
    'This Software Requirements Specification (SRS) document provides a comprehensive description of the '
    'HEALTH AI platform — a European HealthTech co-creation and innovation platform. It defines all '
    'functional, non-functional, and interface requirements for the production release of the system, '
    'and serves as the authoritative reference for development, verification, validation, and stakeholder '
    'review activities. The document follows IEEE 830-1998 and ISO/IEC/IEEE 29148:2018 conventions and '
    'is intended to be the contractual baseline against which the implemented system is evaluated.'
)

doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph(
    'HEALTH AI is a secure, GDPR-compliant single-page web application that enables structured partner '
    'discovery between healthcare professionals and engineers/developers across European institutions. '
    'The platform facilitates cross-disciplinary collaboration through a five-stage workflow: (1) users '
    'register with an institutional .edu address and a one-time email verification code; (2) they post '
    'announcements describing a project, the technical needs and the type of partner sought; (3) other '
    'users browse and filter the feed and open detail pages; (4) interested users accept a non-disclosure '
    'agreement and express interest, optionally proposing a meeting time slot; (5) the post author '
    'accepts a slot, after which the parties move to an external meeting platform (Zoom or Microsoft '
    'Teams) for the actual call. The HEALTH AI platform never stores patient data, intellectual property '
    'documents or meeting recordings — it acts strictly as a neutral first-contact intermediary.'
)
doc.add_paragraph('The current release delivers the following key capabilities:')
items = [
    'User registration and authentication restricted to institutional .edu email addresses',
    'Six-digit OTP email verification (10-minute TTL, 60-second resend cooldown) via EmailJS',
    'SHA-256 client-side password hashing with the Web Crypto API',
    'Three-step Create Announcement wizard (Core Details → Technical Info → Settings) with a visual progress indicator',
    'Innovator Feed dashboard with real-time text search and filters by domain, stage, country, city and status',
    '3D parallax tilt cards with match indicators (same city, complementary role)',
    'Two-column post detail view with NDA-gated technical blueprint and a workflow action sidebar',
    'NDA acceptance modal as the gate before any expression of interest',
    'Auto-generated weekday meeting time slots (10:00, 14:00, 16:00 over the next five business days)',
    'Real-time Firebase chat with typing indicators, single/double check read receipts, reactions and replies',
    'Bookmark / saved-projects tab on the dashboard',
    'Real-time notifications with bell dropdown and system-wide toast alerts',
    'GDPR data export (Article 20) and account deletion (Article 17)',
    'Admin dashboard with overview, user freeze/unfreeze, post moderation and audit log CSV export',
    'Cmd/Ctrl+K Command Palette for global navigation and a Shift+? keyboard shortcuts reference modal',
    'Cinematic editorial landing page with Spline 3D background, scrollytelling, BigTextReveal and BentoFeatures',
    'Responsive design from 320 px mobile to 2560 px desktop, with prefers-reduced-motion support',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'The product\'s primary objectives are to (a) shorten the time to first contact between clinicians '
    'and engineers, (b) provide a regulatory-clean channel for that first contact, and (c) preserve the '
    'confidentiality of pre-IP discussion through an enforced NDA gate.'
)

doc.add_heading('1.3 Definitions, Acronyms, and Abbreviations', level=2)
add_table(
    ['Term', 'Definition'],
    [
        ('SRS', 'Software Requirements Specification'),
        ('SDD', 'Software Design Document'),
        ('GDPR', 'General Data Protection Regulation (EU 2016/679)'),
        ('NDA', 'Non-Disclosure Agreement'),
        ('IP', 'Intellectual Property'),
        ('OTP', 'One-Time Password — six-digit code emailed during registration'),
        ('API', 'Application Programming Interface'),
        ('CRUD', 'Create, Read, Update, Delete'),
        ('SPA', 'Single Page Application'),
        ('CSR', 'Client-Side Rendered'),
        ('RBAC', 'Role-Based Access Control'),
        ('HCP', 'Healthcare Professional'),
        ('UML', 'Unified Modeling Language'),
        ('DFD', 'Data Flow Diagram'),
        ('ERD', 'Entity-Relationship Diagram'),
        ('FR', 'Functional Requirement'),
        ('NFR', 'Non-Functional Requirement'),
        ('TTL', 'Time-To-Live'),
        ('CSR-SPA', 'Client-Side Rendered Single Page Application'),
        ('HMR', 'Hot Module Replacement (Vite dev server)'),
    ],
)

doc.add_heading('1.4 References', level=2)
refs = [
    'IEEE Std 830-1998 — Recommended Practice for Software Requirements Specifications',
    'ISO/IEC/IEEE 29148:2018 — Systems and software engineering — Life cycle processes — Requirements engineering',
    'GDPR — Regulation (EU) 2016/679 of the European Parliament and of the Council',
    'OWASP Top 10 (2024) — https://owasp.org/Top10',
    'WCAG 2.1 — Web Content Accessibility Guidelines',
    'Firebase Firestore Documentation — https://firebase.google.com/docs/firestore',
    'React 19 Documentation — https://react.dev',
    'Framer Motion Documentation — https://www.framer.com/motion',
    'Vite Build Tool — https://vitejs.dev',
    'EmailJS Service — https://www.emailjs.com',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('1.5 Document Conventions', level=2)
doc.add_paragraph(
    'Functional requirements are identified as FR-NN, non-functional requirements as NFR-NN, and use '
    'cases as UC-NN. Each requirement is assigned a priority of High, Medium or Low. The keywords '
    '"shall", "should" and "may" follow RFC 2119 conventions: "shall" denotes mandatory behaviour, '
    '"should" denotes strong recommendation, and "may" denotes optional behaviour. UI element names '
    'are written in Title Case (e.g., "Express Interest" button) and code identifiers in monospace (e.g., '
    'useAuth, addInterestToSubcol).'
)

doc.add_heading('1.6 Intended Audience', level=2)
add_table(
    ['Audience', 'Use of this document'],
    [
        ('Project Sponsors / Course Instructor', 'Verify the project scope and grade against the agreed deliverables'),
        ('Development Team (Health Shield)', 'Implement, refactor and extend features against a single source of truth'),
        ('QA / Reviewer', 'Drive black-box and acceptance test cases from FR/NFR identifiers'),
        ('Future Maintainer', 'Onboard quickly to the existing system and its constraints'),
        ('Compliance Reviewer', 'Confirm GDPR, NDA and audit logging obligations are satisfied'),
    ],
)

doc.add_heading('1.7 Overview', level=2)
doc.add_paragraph(
    'Section 2 gives an overall description of the product, its operating environment, and its '
    'constraints. Section 3 captures use cases and the primary actors. Section 4 enumerates the '
    'functional, non-functional and interface requirements. Section 5 contains the system models — '
    'data flow, entity-relationship and state transition diagrams. Section 6 outlines the system '
    'architecture and technology stack. Section 7 documents the data model in detail. Section 8 '
    'provides a traceability matrix linking requirements to use cases and verification methods. '
    'Section 9 lists acceptance criteria. Appendices A and B contain UI screenshots and a glossary.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. OVERALL DESCRIPTION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. Overall Description', level=1)

doc.add_heading('2.1 Product Perspective', level=2)
doc.add_paragraph(
    'HEALTH AI is a self-contained web application built as a React single page application with '
    'Firebase Firestore as the real-time backend. It does not integrate with hospital information '
    'systems, electronic health records, university research platforms or patent management systems. '
    'It is positioned as a neutral first-contact intermediary that provides a regulator-friendly '
    'discovery surface for healthcare innovators and engineers. The platform is deployed as a static '
    'SPA bundle behind a CDN; all real-time data synchronisation happens through Firestore '
    'onSnapshot listeners over a WebSocket channel from the browser.'
)
doc.add_paragraph(
    'The platform follows a four-layer architecture (Presentation → Business Logic → Data Access → '
    'Infrastructure) and a service-oriented integration with Firebase Firestore (real-time NoSQL), '
    'EmailJS (transactional emails) and external meeting hosts (Zoom/Teams) for the actual video call.'
)

doc.add_heading('2.2 System Context Diagram', level=2)
doc.add_paragraph(
    'The diagram below identifies the actors and external systems that interact with the HEALTH AI '
    'platform and the nature of each interaction.'
)
add_diagram('context_diagram.png', 'System Context — actors and external systems around HEALTH AI')

doc.add_heading('2.3 Product Functions', level=2)
doc.add_paragraph('At a high level, the system shall provide the following major functions:')
functions = [
    ('Authentication & Registration', 'Institutional .edu email-based registration with SHA-256 password hashing and a 6-digit OTP email verification step. Login, logout and a 30-minute inactivity timeout. Personal email providers (Gmail, Yahoo, Outlook, etc.) are blocked.'),
    ('Announcement Management', 'Three-step Create Announcement wizard with title, domain, project stage, description, technical blueprint, required expertise, collaboration type, confidentiality level, location and expiry. Posts support edit, status transitions (Active → Meeting Scheduled → CLOSED) and configurable expiry.'),
    ('Discovery & Filtering', 'Real-time dashboard search across titles, descriptions, domains and expertise fields with multi-dimensional filters by domain, stage, country, city and status, plus All Channels and Saved Projects tabs.'),
    ('Interest & Meeting Workflow', 'NDA acceptance modal before any expression of interest. Auto-generated weekday time slots (10:00, 14:00, 16:00 across five business days). The post author can accept or decline a request; acceptance moves the post to Meeting Scheduled.'),
    ('Real-Time Messaging', 'Firebase-backed 1-on-1 chat with typing indicators, single/double check read receipts, message reactions, reply threading and conversation management (clear history, delete).'),
    ('Profile & GDPR', 'User profile with editable fields, post statistics, GDPR data export (Article 20) and account deletion (Article 17) with confirmation modal.'),
    ('Admin Console', 'Overview cards, user management table with freeze/unfreeze, post moderation table with delete, and audit-log viewer with role/action filters and CSV export.'),
    ('Notifications', 'Real-time bell dropdown with count badge plus system-wide toast alerts for interest expressions, meeting requests and responses.'),
    ('UI / UX shell', 'Cinematic editorial landing page (Spline 3D, BigTextReveal, StickyShowcase, HowItWorks, TwoSides, BentoFeatures), animated page transitions, 3D parallax tilt cards, HLS video backgrounds, skeleton loaders, network status indicator, global error boundary, Cmd+K command palette, Shift+? shortcuts modal, SVG noise grain overlay.'),
]
for ttl, desc in functions:
    p = doc.add_paragraph()
    run = p.add_run(f'{ttl}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.4 User Classes and Characteristics', level=2)
add_table(
    ['Class', 'Profile', 'Frequency', 'Capabilities'],
    [
        ('Healthcare Professional', 'Doctors, clinicians and clinical researchers with domain knowledge and innovation ideas. Comfortable with web tools but typically not software developers.', 'Weekly', 'Create / manage / close posts, browse engineer posts, NDA + interest, propose meetings, chat, profile, GDPR export/delete'),
        ('Engineer / Developer', 'Software, hardware and biomedical engineers. Familiar with technical detail, looking for clinical context and validation partners.', 'Weekly', 'Create / manage / close posts, browse clinician posts, NDA + interest, propose meetings, chat, profile, GDPR export/delete'),
        ('Admin', 'Platform administrator (single super-user account today; auto-assigned for the admin@healthai.edu email). Responsible for moderation and compliance reporting.', 'Daily', 'All user capabilities + freeze/unfreeze users, delete any post, view and export audit logs as CSV'),
    ],
)

doc.add_heading('2.5 Operating Environment', level=2)
add_table(
    ['Aspect', 'Requirement'],
    [
        ('Client browsers', 'Chrome, Firefox, Safari, Edge — latest two major versions'),
        ('Viewport range', '320 px (mobile) to 2560 px (desktop), responsive layouts'),
        ('Network', 'Persistent broadband or 4G/5G connection required for real-time features'),
        ('JavaScript', 'Modern ES2020+, WebSockets, IndexedDB, Web Crypto API'),
        ('Backend', 'Firebase Firestore (Google Cloud), Firebase Auth gateway, EmailJS REST API'),
        ('Hosting', 'Static asset host (e.g., Firebase Hosting, Netlify or Vercel) over HTTPS'),
        ('Build', 'Node.js 20+, npm, Vite 7.3 build pipeline'),
    ],
)

doc.add_heading('2.6 Design and Implementation Constraints', level=2)
constraints = [
    'Only institutional .edu email addresses are accepted at registration; all common personal providers are blocked.',
    'No patient data, medical records or clinical datasets shall be stored on the platform.',
    'No intellectual property documents shall be uploaded — the platform facilitates contact only.',
    'Meetings take place on external platforms (Zoom or Teams); the platform does not host video calls or recordings.',
    'Passwords are hashed client-side with SHA-256 via the Web Crypto API and never stored or transmitted in plaintext.',
    'All persistence is performed through Firebase Firestore; no custom server is operated.',
    'Firestore document size (1 MB) and read/write rate limits constrain the data model.',
    'The platform requires an active internet connection — there is no offline mode beyond a status banner.',
    'All UI animations shall honour the prefers-reduced-motion media query through Framer Motion\'s MotionConfig.',
    'The application shall comply with GDPR data minimisation, right of access, right to portability and right to erasure.',
]
for c in constraints:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('2.7 Assumptions and Dependencies', level=2)
assumptions = [
    'Users have access to a modern browser with JavaScript and WebSockets enabled.',
    'Firebase Firestore remains available with the current SDK API surface (v12).',
    'The EmailJS account remains within the free-tier quota for OTP and meeting notification emails.',
    'Users possess valid institutional .edu email addresses and can receive external email.',
    'External meeting platforms (Zoom / Microsoft Teams) are available to scheduled participants.',
    'Network connectivity is available for the duration of any real-time interaction.',
]
for a in assumptions:
    doc.add_paragraph(a, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. USE CASES
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Use Cases', level=1)

doc.add_heading('3.1 Use Case Diagram', level=2)
doc.add_paragraph(
    'The following UML use case diagram identifies the primary actors (Healthcare Professional, '
    'Engineer/Developer and Admin) and the use cases each may invoke inside the system boundary. '
    'The Admin actor inherits all standard user capabilities (shown by dotted lines) and additionally '
    'has access to administrative use cases.'
)
add_diagram('use_case_diagram.png', 'UML Use Case Diagram — primary actors and system boundary')

doc.add_heading('3.2 Use Case Specifications', level=2)

uc_specs = [
    ('UC-01 Register & Verify OTP',
     'Healthcare Professional or Engineer',
     'A new user creates an account with an institutional .edu email address and verifies it via a 6-digit OTP code received by email.',
     ['User opens the Login page and switches to the Register tab.',
      'User submits all required fields (email, password, name, role, institution, country, city).',
      'System validates the .edu email and that no account already exists for it.',
      'System hashes the password (SHA-256) and dispatches a 6-digit OTP via EmailJS (TTL 10 min, resend cooldown 60 s).',
      'User enters the OTP within the inline panel.',
      'System creates the user document in Firestore and signs the user in.'],
     'User account exists in the users collection and the user is authenticated and redirected to the Dashboard.'),

    ('UC-02 Sign In',
     'Registered user',
     'A user authenticates with their institutional email and password.',
     ['User submits email and password.',
      'System hashes the password and compares against the stored hash.',
      'System writes a LOGIN entry to activityLogs.',
      'System loads the session and redirects to the Dashboard.'],
     'User is authenticated; session is stored in localStorage with a 30-minute inactivity timeout.'),

    ('UC-03 Create Announcement',
     'Healthcare Professional or Engineer',
     'An author publishes a project announcement using the three-step wizard.',
     ['User opens /create-post.',
      'Step 1: enters title, domain, project stage and description.',
      'Step 2: enters technical blueprint (NDA-gated), required expertise, collaboration type.',
      'Step 3: enters location (country/city), confidentiality level and expiry days.',
      'User clicks "Publish Announcement".',
      'System persists the post in Firestore and broadcasts the snapshot to all subscribed clients.'],
     'Post is visible on the Dashboard with status Active.'),

    ('UC-04 Express Interest with NDA',
     'Cross-role user (HCP viewing engineer post or vice versa)',
     'A viewer accepts the NDA, optionally writes a message and registers their interest in the post.',
     ['User opens the post detail page.',
      'User clicks "Express Interest"; the NDA modal appears.',
      'User reads and checks the NDA acknowledgement, optionally writes a message and confirms.',
      'System writes the interest record to posts/{id}/interests inside a transaction that increments interestCount.',
      'System dispatches a notification to the post author.'],
     'Interest is recorded; technical blueprint is unlocked for the viewer if confidentiality is meeting-only.'),

    ('UC-05 Propose & Confirm Meeting',
     'Interested user (proposer) and post author (responder)',
     'After expressing interest the user proposes a time slot and the author accepts or declines.',
     ['Proposer picks a slot from the auto-generated weekday list (next 5 business days × 10:00, 14:00, 16:00).',
      'System writes the meeting request to posts/{id}/meetings and notifies the author.',
      'Author opens the post detail page and clicks Accept or Decline.',
      'On accept: post status changes to Meeting Scheduled and both parties are notified with the external meeting reminder (Zoom/Teams).'],
     'Both users have a Meeting Scheduled notification with a slot and external host instructions.'),

    ('UC-06 Real-Time Chat',
     'Any authenticated user',
     'Two users exchange messages in real time, with typing indicators and read receipts.',
     ['User A opens /chat or clicks "Message" on a post detail page.',
      'System creates or retrieves the conversation document keyed by sorted member IDs.',
      'User A types — system updates conversation.isTyping; user B sees animated dots.',
      'User A sends — system writes a message document; both clients receive the snapshot.',
      'User B opens the conversation; system marks unread messages as read; user A sees a double-check.'],
     'Both users have an up-to-date conversation thread with delivery and read state synchronised.'),

    ('UC-07 GDPR Data Export and Account Deletion',
     'Any authenticated user',
     'A user exercises GDPR rights to download their data or delete their account.',
     ['User opens the Profile page → Data Privacy section.',
      'For export: user clicks "Export My Data (JSON)" and the browser downloads a JSON file containing the profile and all posts.',
      'For deletion: user clicks "Delete My Account (Art. 17)", confirms in the modal; system removes the user document and cascades the deletion.'],
     'Either: a JSON file is downloaded; or the user account and personal data are removed from Firestore.'),

    ('UC-08 Admin Audit Review',
     'Admin',
     'The administrator reviews user activity and exports a CSV for compliance reporting.',
     ['Admin opens /admin and selects the Audit Logs tab.',
      'Admin filters by role, action type or text and inspects matching events.',
      'Admin clicks "Export CSV"; the browser downloads the filtered audit log.'],
     'Admin obtains a CSV file of audit log entries matching the chosen filters.'),
]

for name, actor, desc, steps, post in uc_specs:
    doc.add_heading(name, level=3)
    p = doc.add_paragraph()
    p.add_run('Primary Actor: ').bold = True
    p.add_run(actor)
    p = doc.add_paragraph()
    p.add_run('Description: ').bold = True
    p.add_run(desc)
    p = doc.add_paragraph()
    p.add_run('Main Flow:').bold = True
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')
    p = doc.add_paragraph()
    p.add_run('Postcondition: ').bold = True
    p.add_run(post)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. SPECIFIC REQUIREMENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. Specific Requirements', level=1)

doc.add_heading('4.1 Functional Requirements', level=2)
fr_data = [
    ('FR-01', 'User Registration', 'The system shall allow users to register with .edu email, password (min 6 chars), full name with title, role, institution, country and city. Personal providers are blocked.', 'High'),
    ('FR-02', 'OTP Verification', 'The system shall send a 6-digit OTP via EmailJS during registration with a 10-minute TTL and 60-second resend cooldown.', 'High'),
    ('FR-03', 'User Login', 'Registered users shall log in with email and SHA-256-hashed password. Successful and failed attempts shall be logged.', 'High'),
    ('FR-04', 'User Logout', 'Authenticated users shall be able to log out from the navbar. Session is cleared from local state and storage.', 'High'),
    ('FR-05', 'Inactivity Timeout', 'The system shall log the user out automatically after 30 minutes of inactivity (no mouse/keyboard/scroll/touch events).', 'Medium'),
    ('FR-06', 'Admin Auto-Detection', 'Registration with admin@healthai.edu shall auto-assign the Admin role.', 'Medium'),
    ('FR-07', 'Create Announcement Wizard', 'The system shall provide a 3-step wizard with progress indicator: Core Details → Technical Info → Settings.', 'High'),
    ('FR-08', 'Browse Announcements', 'The dashboard shall display all non-deleted posts with search, domain, stage, country, city and status filters; default filter shows Active posts.', 'High'),
    ('FR-09', 'Post Detail View', 'The system shall present each post in a two-column layout: main content (status, title, metadata, overview, blueprint, expertise) and sidebar (counterparty, workflow actions, author card).', 'High'),
    ('FR-10', 'Edit Post', 'Authors shall be able to edit title and description through a modal dialog.', 'Medium'),
    ('FR-11', 'Close Post', 'Authors shall be able to mark a post as Partner Found (CLOSED).', 'High'),
    ('FR-12', 'Bookmark Posts', 'Users shall be able to bookmark posts; bookmarked posts appear in a Saved Projects tab.', 'Medium'),
    ('FR-13', 'Post Expiry', 'Each post shall have a configurable expiry of 15, 30, 60 or 90 days. Expired posts shall be flagged automatically.', 'Low'),
    ('FR-14', 'Express Interest', 'Cross-role users shall be able to express interest after explicit NDA acceptance, with an optional message.', 'High'),
    ('FR-15', 'NDA Modal', 'Interest expression shall require explicit acknowledgement of NDA terms via a checkbox modal.', 'High'),
    ('FR-16', 'Propose Meeting', 'After expressing interest, the user shall see auto-generated weekday slots for the next 5 business days at 10:00, 14:00 and 16:00 and may select one.', 'High'),
    ('FR-17', 'Respond to Meeting', 'The post author shall be able to accept or decline a meeting request. Acceptance transitions the post to Meeting Scheduled and notifies both parties.', 'High'),
    ('FR-18', 'Confidential Posts', 'Posts with meeting-only confidentiality shall hide the Technical Blueprint section for non-authors until the viewer signs the NDA.', 'High'),
    ('FR-19', 'Real-Time Messaging', 'The system shall provide Firebase-backed 1-on-1 chat with conversation list, message history and real-time updates.', 'High'),
    ('FR-20', 'Typing Indicators', 'The system shall surface a real-time typing indicator to the other participant via Firestore field updates.', 'Medium'),
    ('FR-21', 'Read Receipts', 'Messages shall display single check (sent) and double check (read) indicators.', 'Medium'),
    ('FR-22', 'Message Reactions & Replies', 'The system shall support emoji reactions on messages and threaded replies.', 'Medium'),
    ('FR-23', 'User Discovery in Chat', 'The chat sidebar shall let users discover and start conversations with any registered user.', 'Medium'),
    ('FR-24', 'Chat Management', 'Users shall be able to clear message history or delete entire conversations.', 'Low'),
    ('FR-25', 'Message from Post', 'The post detail page shall expose a "Message" action that opens a chat with the post author.', 'Medium'),
    ('FR-26', 'Edit Profile', 'Users shall be able to edit name, institution, country and city through an inline form.', 'Medium'),
    ('FR-27', 'GDPR Data Export', 'Users shall be able to export all their data as a JSON file (GDPR Art. 20).', 'High'),
    ('FR-28', 'GDPR Account Deletion', 'Users shall be able to permanently delete their account through a confirmation modal (GDPR Art. 17).', 'High'),
    ('FR-29', 'Admin Overview', 'The admin dashboard shall present overview cards with total users, total announcements, active projects and monitored events.', 'Medium'),
    ('FR-30', 'Admin User Management', 'Admins shall be able to view all users in a table and freeze or unfreeze accounts.', 'High'),
    ('FR-31', 'Admin Post Moderation', 'Admins shall be able to view all posts and delete any post.', 'High'),
    ('FR-32', 'Audit Log Viewer', 'The admin shall be able to filter, search and export audit logs as CSV.', 'High'),
    ('FR-33', 'Notification Bell', 'The system shall surface real-time notifications for interest, meeting requests and responses through a bell dropdown with count badge.', 'Medium'),
    ('FR-34', 'Notification Dismissal', 'Users shall be able to dismiss notifications individually or all at once.', 'Low'),
    ('FR-35', 'Toast Notifications', 'The system shall show transient toast alerts for critical events outside the bell dropdown.', 'Medium'),
    ('FR-36', 'Command Palette', 'The system shall provide a global Cmd/Ctrl+K command palette for navigation and search.', 'Medium'),
    ('FR-37', 'Keyboard Shortcuts Modal', 'Pressing Shift+? shall open a reference modal listing all available keyboard bindings.', 'Low'),
    ('FR-38', 'Page Transitions', 'The system shall use animated page transitions through Framer Motion AnimatePresence.', 'Low'),
    ('FR-39', 'Match Indicators', 'Dashboard cards shall show match indicators for same-city or complementary-role posts.', 'Medium'),
    ('FR-40', 'Responsive Design', 'The application shall adapt from 320 px to 2560 px viewport widths.', 'Medium'),
    ('FR-41', 'Reduced Motion', 'The system shall honour prefers-reduced-motion through Framer Motion MotionConfig.', 'Medium'),
    ('FR-42', 'Network Status Indicator', 'The system shall display an offline banner when navigator.onLine is false.', 'Medium'),
    ('FR-43', 'Global Error Boundary', 'A React Error Boundary shall catch rendering failures and present a recovery UI rather than a blank screen.', 'High'),
    ('FR-44', 'Activity Logging', 'Login, login failure, post creation, post close, meeting request and admin actions shall be persisted to the activityLogs collection.', 'High'),
    ('FR-45', 'Online Presence', 'The system shall track user online state (isOnline, lastSeen) using a Firestore presence field.', 'Low'),
]

t = doc.add_table(rows=1, cols=4)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['ID', 'Name', 'Description', 'Priority']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
for fr_id, name, desc, priority in fr_data:
    row = t.add_row()
    row.cells[0].text = fr_id
    row.cells[1].text = name
    row.cells[2].text = desc
    row.cells[3].text = priority
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(8)
doc.add_paragraph()

doc.add_heading('4.2 Non-Functional Requirements', level=2)
nfrs = [
    ('NFR-01', 'Performance', 'Pages shall load to interactive within 2 s on a standard 10 Mbps broadband connection. Real-time updates (messages, notifications) shall appear within 500 ms of the originating event.'),
    ('NFR-02', 'Security — Hashing', 'Passwords shall be hashed client-side with SHA-256 via Web Crypto API. No plaintext password shall ever leave the browser.'),
    ('NFR-03', 'Security — Email Domain', 'Only institutional .edu email addresses shall be accepted at registration; common personal providers are blocked.'),
    ('NFR-04', 'Security — Activity Logging', 'All authentication events (success and failure), post mutations and admin actions shall be persisted to an immutable audit log.'),
    ('NFR-05', 'Security — RBAC', 'Admin-only routes shall be guarded; non-admin users attempting to access /admin shall be redirected.'),
    ('NFR-06', 'Privacy — GDPR Art. 15', 'Users shall be able to view all personal data the system holds about them on the Profile page.'),
    ('NFR-07', 'Privacy — GDPR Art. 20', 'Users shall be able to export their data as a machine-readable JSON file.'),
    ('NFR-08', 'Privacy — GDPR Art. 17', 'Users shall be able to delete their account and all associated personal data permanently.'),
    ('NFR-09', 'Privacy — Data Minimisation', 'The system shall not store patient data, medical records, IP documents or meeting recordings.'),
    ('NFR-10', 'Scalability', 'Firebase Firestore shall scale horizontally to support up to ~10 000 concurrent connected clients without changes to the application code.'),
    ('NFR-11', 'Usability', 'Any feature shall be reachable within three clicks from the dashboard. The Cmd+K palette shall offer single-shortcut navigation to any page.'),
    ('NFR-12', 'Reliability', 'A global React Error Boundary shall catch rendering failures and present a recovery UI; transient Firestore failures shall be surfaced through toast notifications without crashing the page.'),
    ('NFR-13', 'Availability', 'The hosted SPA shall achieve at least 99% monthly uptime, contingent on the underlying Firebase availability SLA.'),
    ('NFR-14', 'Accessibility', 'The UI shall use semantic HTML5, accessible names on interactive elements, sufficient colour contrast (WCAG AA) and shall honour prefers-reduced-motion.'),
    ('NFR-15', 'Maintainability', 'Business logic shall live in custom hooks (useAuth, usePosts, useChat, useNotifications, usePostEngagement) and persistence in services/firestore.js to keep components focused on rendering.'),
    ('NFR-16', 'Compatibility', 'The application shall function on the latest two major versions of Chrome, Firefox, Safari and Edge.'),
    ('NFR-17', 'Localisation Readiness', 'All UI strings are extractable from JSX so that future localisation can be added without rewriting components.'),
    ('NFR-18', 'Observability', 'Critical errors shall log to console.error and surface a user-facing toast; admin audit logs provide system-level traceability.'),
    ('NFR-19', 'Build & Deployability', 'The system shall build with a single `npm run build` and deploy as a static dist/ to any CDN. Environment configuration shall flow through VITE_FIREBASE_* variables.'),
    ('NFR-20', 'Compliance — Email Verification', 'No user account shall become active until OTP verification has succeeded.'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['ID', 'Category', 'Requirement']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
for nfr_id, cat, req in nfrs:
    row = t.add_row()
    row.cells[0].text = nfr_id
    row.cells[1].text = cat
    row.cells[2].text = req
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
doc.add_paragraph()

doc.add_heading('4.3 External Interface Requirements', level=2)
doc.add_heading('4.3.1 User Interface', level=3)
doc.add_paragraph(
    'The user interface shall follow the editorial liquid-glass design system documented in Section 4.4 '
    'and shall remain visually consistent across all routes. All interactive elements shall be reachable '
    'by keyboard.'
)

doc.add_heading('4.3.2 Software Interfaces', level=3)
add_table(
    ['External Interface', 'Purpose', 'Protocol'],
    [
        ('Firebase Firestore SDK v12', 'Persistence and real-time subscriptions', 'WebSocket / HTTPS'),
        ('EmailJS REST API', 'Send 6-digit OTP and meeting notification emails', 'HTTPS POST'),
        ('Zoom / Microsoft Teams', 'External meeting host links surfaced after meeting acceptance', 'HTTPS link'),
        ('Spline runtime', '3D landing-page background scene', 'HTTPS / iframe'),
    ],
)

doc.add_heading('4.3.3 Hardware Interfaces', level=3)
doc.add_paragraph(
    'The system has no direct hardware dependency beyond a standard internet-connected device with a '
    'modern web browser. WebGL support is required for the 3D landing background; the application '
    'gracefully degrades to a static image where WebGL is unavailable.'
)

doc.add_heading('4.3.4 Communication Interfaces', level=3)
doc.add_paragraph(
    'All client-server traffic uses HTTPS. Real-time data is exchanged through the Firestore SDK over '
    'a WebSocket channel. Email is sent through the EmailJS HTTPS API; the system never sends mail '
    'directly. There is no inbound API to the SPA.'
)

doc.add_heading('4.4 User Interface Requirements', level=2)
doc.add_paragraph(
    'The HEALTH AI platform features a premium dark-mode interface with glassmorphism design elements, '
    'gradient accents, and smooth micro-animations. The following screenshots represent the current '
    'production state of the application as of April 2026.'
)

ui_screens = [
    ('3.3.1 Landing Page',
     'The landing page uses a cinematic editorial scrollytelling flow with a persistent Spline 3D '
     'interactive background. The sequence is: (1) Hero with gradient text and CTA buttons, '
     '(2) BigTextReveal with per-line clip-path mask animations, (3) StickyShowcase — a 3-act pinned '
     'scroll experience (160 vh) presenting platform value propositions with SVG mockup illustrations, '
     '(4) HowItWorks — 4-step workflow (Post → Discovery → NDA → Meeting) with an animated SVG '
     'connector path, (5) TwoSides split-pane persona view, (6) BentoFeatures — 4 capability cards. '
     'A 3.5% SVG noise grain overlay adds premium texture. All animations respect prefers-reduced-motion.',
     'landing_page.png',
     'Landing Page — Cinematic editorial flow with Spline 3D background and scrollytelling sections'),
    ('3.3.2 Login / Registration Page',
     'The authentication page supports both Sign In and Register modes via an animated segmented '
     'tab switcher. After Register submission, a 6-digit OTP verification panel appears where users '
     'enter the code sent through EmailJS. Features include inline field icons, animated error '
     'banners, password visibility toggle, gradient logo, and GDPR / NDA trust indicators.',
     'login_page.png',
     'Login / Registration — Liquid-glass editorial panel with .edu validation and OTP'),
    ('3.3.3 Dashboard (Innovator Feed)',
     'The main dashboard displays all announcements in a responsive grid layout with 3D parallax tilt '
     'cards. The header shows "Innovator Feed" with a "New Announcement" button. Tabs switch between '
     'All Channels and Saved Projects. Filters include real-time text search, domain, stage, country, '
     'city and status with a result counter. Each card shows role badge, domain badge, status badge, '
     'title, description excerpt, match indicators, required partner type, author avatar, and bookmark.',
     'dashboard_page.png',
     'Dashboard — Innovator Feed with multi-filter announcement grid'),
    ('3.3.4 Create Announcement (Multi-Step Wizard)',
     'The Create Announcement form uses a 3-step wizard with visual progress indicator. Step 1 '
     'gathers Core Details (title, domain selector, project stage cards, description). Step 2 '
     'gathers Technical Info (blueprint, required expertise, collaboration type). Step 3 gathers '
     'Settings (country/city, confidentiality cards Public Info / NDA Protected, expiry).',
     'create_post_page.png',
     'Create Announcement — Step 1: Core Details with project stage cards'),
    ('3.3.5 Post Detail Page',
     'Two-column layout: main content panel with status badges, post title, metadata chips '
     '(stage, location, date, expiry), Executive Overview section, Technical Blueprint (locked when '
     'confidential), and Required Expertise section. The sidebar shows the target counterparty card, '
     'workflow action buttons (Express Interest → Propose Meeting → Meeting Confirmed), and a '
     '"Posted By" author card with a direct message button.',
     'post_detail_page.png',
     'Post Detail — Two-column layout with NDA-gated technical blueprint'),
    ('3.3.6 Chat / Messages Page',
     'Two-panel layout: left sidebar with the conversation list (search, avatars, last messages, '
     'timestamps, unread badges, delete options) and a New Chat user discovery list. The main chat '
     'area shows date-grouped messages, sender-specific bubble alignment, animated typing dots, '
     'single/double-check read receipts and a message input with conditional send button. An options '
     'menu provides Clear History and Delete Conversation actions.',
     'chat_page.png',
     'Messages — Real-time chat with conversation sidebar and read receipts'),
    ('3.3.7 Profile Page',
     'Two-column layout: left card with user avatar (gradient initial), name, role badge, email, '
     'institution, location, join date and post statistics; right panel with Profile Settings '
     '(editable fields with edit/save/cancel) and a Data Privacy (GDPR) section with "Export My Data '
     '(JSON)" and "Delete My Account (Art. 17)" buttons.',
     'profile_page.png',
     'Profile — User info, settings and GDPR data privacy controls'),
]
for h, body, scr, cap in ui_screens:
    doc.add_heading(h, level=3)
    doc.add_paragraph(body)
    add_screenshot(scr, cap)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. SYSTEM MODELS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. System Models', level=1)
doc.add_paragraph(
    'This section provides the formal models that describe how data flows through the system, how '
    'persistent entities relate to each other, and how the announcement lifecycle progresses through '
    'its discrete states.'
)

doc.add_heading('5.1 Data Flow Diagram (Level 1)', level=2)
doc.add_paragraph(
    'The Level 1 DFD below shows the major processes in the system, the data stores they read from '
    'and write to, and the external entities that exchange data with the platform. P1–P9 correspond '
    'to the major functional areas, while D1–D6 correspond to Firestore collections.'
)
add_diagram('dfd_level1.png', 'Data Flow Diagram (Level 1) — processes, stores and external entities')

doc.add_heading('5.2 Entity-Relationship Diagram', level=2)
doc.add_paragraph(
    'The ERD below documents the persistent entities of HEALTH AI as represented in Firestore. '
    'Subcollections (posts/{id}/interests, posts/{id}/meetings, conversations/{id}/messages) are shown '
    'as separate entities with their owning collection prefixed in the title.'
)
add_diagram('er_diagram.png', 'Entity-Relationship Diagram — Firestore data model with cardinalities')

doc.add_heading('5.3 State Transition Diagram', level=2)
doc.add_paragraph(
    'The diagram below documents the lifecycle of a post — the central domain entity. A post starts '
    'in Active when published and transitions through Active + Interest Received and Meeting '
    'Scheduled before reaching the terminal CLOSED (Partner Found) state. Two side-paths exist for '
    'expiry and admin deletion.'
)
add_diagram('state_diagram.png', 'Post Lifecycle — UML state transition diagram')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. System Architecture', level=1)

doc.add_heading('6.1 Architectural Overview', level=2)
doc.add_paragraph(
    'HEALTH AI follows a four-layer client-side rendered architecture with Firebase as the serverless '
    'backend. The Presentation layer renders React components and animations. The Business Logic '
    'layer encapsulates state and orchestration in custom hooks. The Data Access layer (services/'
    'firestore.js) translates hook calls into Firestore reads, writes and subscriptions. The '
    'Infrastructure layer is provided by Firebase Firestore and the EmailJS service.'
)
add_diagram('architecture_diagram.png', 'Four-layer Architecture — Presentation, Business Logic, Data Access, Infrastructure')

doc.add_heading('6.2 Technology Stack', level=2)
add_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ('Frontend Framework', 'React 19.2 (JSX)', 'Component-based UI rendering with virtual DOM'),
        ('Build Tool', 'Vite 7.3', 'Fast HMR dev server and optimised production bundling'),
        ('Routing', 'React Router v7', 'Client-side routing with ProtectedRoute / GuestRoute guards'),
        ('Animation', 'Framer Motion 12', 'Page transitions, scroll-triggered reveals, MotionConfig for reduced motion'),
        ('3D / WebGL', 'Three.js + React Three Fiber + Drei + Spline', 'Interactive 3D landing background and particle effects'),
        ('Video', 'HLS.js', 'Adaptive-bitrate looping video backgrounds with crossfade'),
        ('Icons', 'Lucide React', 'Tree-shakable SVG icon library'),
        ('Database', 'Firebase Firestore v12', 'Real-time NoSQL document database with subcollections and transactions'),
        ('Email', 'EmailJS', 'OTP verification and meeting-request notifications'),
        ('Authentication', 'Custom (Firestore + SHA-256 + OTP)', 'Client-side auth with hashed passwords and email verification'),
        ('Styling', 'Vanilla CSS + CSS Variables + liquid-glass.css', 'Editorial design system with warm amber / cool cyan palette'),
        ('State Management', 'React Hooks (custom)', 'useAuth, usePosts, useChat, useNotifications, usePostEngagement, useToast'),
        ('Utilities', 'classnames, tailwind-merge', 'Conditional class composition'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. DATA MODEL
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. Data Model', level=1)
doc.add_paragraph('The Firestore database uses the following collections and subcollections.')

doc.add_heading('7.1 Users Collection', level=2)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ('id', 'string', 'Unique user identifier (admin-TS pattern)'),
        ('name', 'string', 'Full name with title (e.g., "Dr. Jane Doe")'),
        ('email', 'string', 'Institutional .edu email (lowercased)'),
        ('passwordHash', 'string', 'SHA-256 hashed password'),
        ('role', 'string', 'Healthcare Professional | Engineer | Admin'),
        ('institution', 'string', 'University or organisation'),
        ('city', 'string', 'User city'),
        ('country', 'string', 'User country'),
        ('registeredAt', 'string (ISO)', 'Registration timestamp'),
        ('status', 'string', 'active | frozen'),
        ('lastLogin', 'string (ISO)', 'Last successful login timestamp'),
        ('isOnline', 'boolean', 'Current online presence'),
        ('lastSeen', 'timestamp', 'Server timestamp of last activity'),
        ('savedPosts', 'array<string>', 'Bookmarked post IDs'),
    ],
)

doc.add_heading('7.2 Posts Collection (with subcollections)', level=2)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ('id', 'string', 'Unique post identifier (post-TS)'),
        ('title', 'string', 'Announcement title'),
        ('domain', 'string', 'Medical / tech domain (cardiology, oncology, etc.)'),
        ('projectStage', 'string', 'idea | concept validation | prototype developed | pilot testing | pre-deployment'),
        ('explanation', 'string', 'Project description (executive overview)'),
        ('highLevelIdea', 'string', 'Technical blueprint (NDA-gated when confidentiality = meeting-only)'),
        ('expertiseNeeded', 'string', 'Required partner expertise'),
        ('collaborationType', 'string', 'co-development | advisory | licensing | joint research | pilot partnership'),
        ('confidentiality', 'string', 'overview-public | meeting-only'),
        ('city / country', 'string', 'Project location'),
        ('authorId / authorName / authorRole / authorEmail', 'string', 'Denormalised author fields'),
        ('createdAt', 'string (ISO)', 'Creation timestamp'),
        ('expiryDate', 'string (ISO)', 'Expiry timestamp'),
        ('status', 'string', 'Active | Meeting Scheduled | CLOSED | Expired | DELETED'),
        ('interestCount', 'number', 'Denormalised count, atomically incremented in transaction'),
        ('meetingCount', 'number', 'Denormalised count, atomically incremented in transaction'),
    ],
)
doc.add_paragraph(
    'posts/{id}/interests subcollection: { userId, userName, message, ndaAccepted, createdAt (serverTimestamp) }.'
)
doc.add_paragraph(
    'posts/{id}/meetings subcollection: { proposedBy, proposedByName, slot (ISO), status (pending|accepted|declined), createdAt (serverTimestamp) }.'
)

doc.add_heading('7.3 Conversations & Messages', level=2)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ('conversations.id', 'string', 'Sorted member IDs joined by underscore'),
        ('members', 'array<string>', 'Two user IDs'),
        ('memberData', 'map', 'Per-user name and role for display'),
        ('lastMessage', 'string', 'Preview of the last message'),
        ('updatedAt', 'timestamp', 'Server timestamp of last activity'),
        ('unreadCount', 'map<userId,int>', 'Per-user unread count'),
        ('isTyping', 'map<userId,bool>', 'Typing indicator field'),
    ],
)
doc.add_paragraph(
    'conversations/{id}/messages subcollection: { senderId, senderName, text, timestamp (serverTimestamp), read, reactions: map<emoji,array<userId>>, replyTo: { messageId, text, senderName } }.'
)

doc.add_heading('7.4 Activity Logs & Notifications', level=2)
add_table(
    ['Collection', 'Key Fields'],
    [
        ('activityLogs', 'id, timestamp, userId, userName, role, actionType, targetEntity, result, details'),
        ('notifications', 'id, userId (recipient), type, title, message, postId, fromUser, timestamp, read'),
    ],
)
doc.add_paragraph(
    'actionType values include LOGIN, LOGIN_FAILED, LOGOUT, POST_CREATE, POST_CLOSE, POST_DELETE, '
    'INTEREST_EXPRESSED, MEETING_REQUEST, MEETING_ACCEPT, MEETING_DECLINE, USER_FREEZE, USER_UNFREEZE, '
    'GDPR_EXPORT, ACCOUNT_DELETE.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. TRACEABILITY MATRIX
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. Requirements Traceability Matrix', level=1)
doc.add_paragraph(
    'The matrix below maps each use case to the functional requirements that realise it and the '
    'verification method used to confirm correct behaviour. This ensures every requirement is '
    'traceable both backwards to a stakeholder need and forwards to a test artefact.'
)
add_table(
    ['Use Case', 'Realising FRs', 'Key NFRs', 'Verification'],
    [
        ('UC-01 Register & Verify OTP', 'FR-01, FR-02, FR-44', 'NFR-02, NFR-03, NFR-20', 'Manual + integration test of registration flow with EmailJS sandbox'),
        ('UC-02 Sign In', 'FR-03, FR-04, FR-05, FR-44', 'NFR-02, NFR-04', 'Manual login with valid/invalid credentials; verify activityLog entry'),
        ('UC-03 Create Announcement', 'FR-07, FR-08, FR-09, FR-13, FR-44', 'NFR-01, NFR-15', 'Manual end-to-end wizard run; check Firestore document shape'),
        ('UC-04 Express Interest with NDA', 'FR-14, FR-15, FR-18, FR-33, FR-44', 'NFR-04, NFR-09', 'Cross-role manual test; verify NDA modal blocks bypass and notification fires'),
        ('UC-05 Propose & Confirm Meeting', 'FR-16, FR-17, FR-33, FR-35, FR-44', 'NFR-01, NFR-04', 'Two-account manual test; check status transition to Meeting Scheduled'),
        ('UC-06 Real-Time Chat', 'FR-19, FR-20, FR-21, FR-22, FR-23, FR-24, FR-25', 'NFR-01, NFR-12', 'Two-browser manual test; verify typing indicator and read receipts'),
        ('UC-07 GDPR Export & Deletion', 'FR-27, FR-28, FR-44', 'NFR-06, NFR-07, NFR-08', 'Manual test; inspect downloaded JSON; verify Firestore documents removed'),
        ('UC-08 Admin Audit Review', 'FR-29, FR-30, FR-31, FR-32, FR-44', 'NFR-04, NFR-05', 'Manual test as admin; verify CSV export contents and freeze enforcement'),
        ('UI shell (cross-cutting)', 'FR-36, FR-37, FR-38, FR-39, FR-40, FR-41, FR-42, FR-43', 'NFR-11, NFR-12, NFR-14, NFR-16', 'Browser matrix smoke test; reduced-motion media query toggle test'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9. ACCEPTANCE CRITERIA
# ══════════════════════════════════════════════════════════════════
doc.add_heading('9. Acceptance Criteria & Verification', level=1)
doc.add_paragraph(
    'The system shall be considered ready for final delivery when all of the following acceptance '
    'criteria are demonstrably satisfied:'
)
ac = [
    ('AC-01', 'Registration with a Gmail/Yahoo/Outlook address is rejected with a clear error message.'),
    ('AC-02', 'Registration with a valid .edu email triggers an OTP email and refuses to create the account until the OTP is verified.'),
    ('AC-03', 'Login with the wrong password produces a LOGIN_FAILED entry in activityLogs and a non-leaking error toast.'),
    ('AC-04', 'A new announcement created by user A appears on user B\'s dashboard within 1 second without a manual reload.'),
    ('AC-05', 'When viewing a meeting-only post, the technical blueprint is hidden until the viewer signs the NDA via the modal.'),
    ('AC-06', 'After the post author accepts a meeting request, both parties see a Meeting Scheduled banner and a notification within 1 second.'),
    ('AC-07', 'Messages between two users render in both browsers within 500 ms; typing indicators and read receipts behave as specified.'),
    ('AC-08', 'GDPR Export downloads a JSON file containing every user field and every post owned by the user.'),
    ('AC-09', 'GDPR Account Deletion removes the user document and the user is signed out; subsequent login fails.'),
    ('AC-10', 'Admin freeze on a user prevents that user from logging in until unfrozen.'),
    ('AC-11', 'CSV export from the admin audit log opens correctly in Excel/Numbers with the displayed columns.'),
    ('AC-12', 'Cmd+K opens the command palette from any authenticated route; Shift+? opens the shortcuts modal.'),
    ('AC-13', 'When prefers-reduced-motion is set, all Framer Motion entrances and exits are skipped.'),
    ('AC-14', 'When the browser goes offline, the network status indicator banner is displayed.'),
    ('AC-15', 'A thrown render error is caught by the GlobalErrorBoundary and the user is offered a recovery button rather than a blank screen.'),
]
add_table(['ID', 'Criterion'], ac)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 10. APPENDIX A — UI SCREENSHOTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('10. Appendix A — UI Screenshots', level=1)
doc.add_paragraph(
    'The following screenshots were captured from the live application on May 1, 2026 and '
    'represent the current state of the HEALTH AI platform user interface after the major UI refactor.'
)

screenshots = [
    ('landing_page.png', 'Landing Page — Cinematic editorial flow with Spline 3D background, BigTextReveal, StickyShowcase scrollytelling, HowItWorks workflow, and TwoSides personas'),
    ('login_page.png', 'Authentication Page — Liquid-glass editorial panel with segmented Sign In / Register tabs, .edu validation, OTP verification, and GDPR trust indicators'),
    ('dashboard_page.png', 'Dashboard (Innovator Feed) — Multi-filter announcement grid with 3D tilt cards, match indicators, role badges, and bookmark functionality'),
    ('create_post_page.png', 'Create Announcement Wizard — 3-step wizard with WizardProgress indicator, domain selection, project stage cards, and settings'),
    ('post_detail_page.png', 'Post Detail — Two-column layout with executive overview, NDA-gated technical blueprint, interest/meeting subcollections, and author card'),
    ('chat_page.png', 'Messages — Real-time chat with conversation sidebar, typing indicators, read receipts, emoji reactions, and message replies'),
    ('profile_page.png', 'Profile — User card with role badge, profile settings editor, post statistics, and GDPR data privacy controls (export & deletion)'),
]
for filename, caption in screenshots:
    add_screenshot(filename, caption, width=Inches(5.6))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 11. APPENDIX B — GLOSSARY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('11. Appendix B — Glossary', level=1)
add_table(
    ['Term', 'Meaning'],
    [
        ('Announcement / Post', 'A user-published project description seeking a complementary partner.'),
        ('Author', 'The user who created a given post.'),
        ('Counterparty', 'The complementary role being sought by the author (engineer for clinician posts and vice versa).'),
        ('Confidentiality', 'A post-level setting controlling whether the technical blueprint is publicly visible or NDA-gated.'),
        ('Express Interest', 'The action of indicating willingness to collaborate on a post, gated behind the NDA modal.'),
        ('Innovator Feed', 'The product name of the dashboard view listing all active announcements.'),
        ('Meeting-only', 'A confidentiality level requiring NDA acceptance before the technical blueprint is revealed.'),
        ('Meeting slot', 'An auto-generated weekday time at 10:00, 14:00 or 16:00 within the next five business days.'),
        ('NDA modal', 'The modal dialog presenting non-disclosure terms before any expression of interest.'),
        ('OTP', 'A six-digit one-time password emailed during registration.'),
        ('Saved Project', 'A bookmarked post; visible under the Saved Projects tab on the dashboard.'),
        ('Toast', 'A transient notification surfaced by the ToastProvider for critical events.'),
        ('Subcollection', 'A Firestore-native collection nested under a parent document (e.g., posts/{id}/interests).'),
    ],
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'HEALTH_AI_SRS_Document.docx')
doc.save(output_path)
print(f'SRS Document saved to: {output_path}')
