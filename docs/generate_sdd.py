#!/usr/bin/env python3
"""Generate HEALTH AI SDD Document — comprehensive final version with UML diagrams."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ───────── PAGE SETUP ─────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

for lv in range(1, 4):
    h = doc.styles[f'Heading {lv}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h.font.bold = True
    h.font.size = Pt([0, 22, 16, 13][lv])
    if lv == 1:
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif lv == 2:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

SDIR = os.path.join(os.path.dirname(__file__), 'screenshots')
DDIR = os.path.join(os.path.dirname(__file__), 'diagrams')


def img_screen(fn, cap, w=Inches(5.6)):
    fp = os.path.join(SDIR, fn)
    if os.path.exists(fp):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(fp, width=w)
        c = doc.add_paragraph(f'Figure: {cap}')
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.italic = True
        c.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def img_diag(fn, cap, w=Inches(6.2)):
    fp = os.path.join(DDIR, fn)
    if os.path.exists(fp):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(fp, width=w)
        c = doc.add_paragraph(f'Diagram: {cap}')
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.italic = True
        c.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()


def tbl(hdrs, rows):
    t = doc.add_table(rows=1, cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(hdrs):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    for rd in rows:
        r = t.add_row()
        for i, tx in enumerate(rd):
            r.cells[i].text = str(tx)
            r.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('HEALTH AI'); r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = RGBColor(0x3b, 0x3b, 0x98)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('European HealthTech Co-Creation & Innovation Platform').font.size = Pt(14)
doc.add_paragraph()
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run('Software Design Document (SDD)'); r.font.size = Pt(20); r.font.bold = True
doc.add_paragraph()
v = doc.add_paragraph(); v.alignment = WD_ALIGN_PARAGRAPH.CENTER
v.add_run('Version 3.0 — Final Release • April 2026').font.size = Pt(12)
doc.add_paragraph()

gp = doc.add_paragraph(); gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
gr = gp.add_run('Group: Health Shield'); gr.font.size = Pt(16); gr.font.bold = True; gr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for m in ['Cem Özal', 'Emre Kurubaş', 'Hasabu Can Eltayeb', 'Sertaç Ataç']:
    mr = mp.add_run(m + '\n'); mr.font.size = Pt(12); mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

it = doc.add_table(rows=0, cols=2); it.style = 'Light Grid Accent 1'
for lb, vl in [('Document', 'Software Design Document'),
               ('Project', 'HEALTH AI Platform'),
               ('Group', 'Health Shield'),
               ('Members', 'Cem Özal, Emre Kurubaş, Hasabu Can Eltayeb, Sertaç Ataç'),
               ('Version', '3.1 (Final)'),
               ('Date', 'May 1, 2026'),
               ('Course', 'SENG 384 — Software Engineering'),
               ('Standard', 'IEEE 1016-2009 Software Design Descriptions'),
               ('Status', 'Final — UML diagrams and module specifications')]:
    rw = it.add_row(); rw.cells[0].text = lb; rw.cells[1].text = vl
    rw.cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# REVISION HISTORY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('Revision History', level=1)
tbl(['Version', 'Date', 'Author(s)', 'Description of Changes'],
    [('1.0', '2025-12-12', 'Health Shield', 'Initial SDD — module decomposition'),
     ('2.0', '2026-04-01', 'Health Shield', 'Refreshed for production build, added screenshots'),
     ('3.0', '2026-04-27', 'Health Shield', 'Final — added Use Case, Architecture, Component, Sequence, Class and Deployment diagrams; expanded module specifications, security & deployment sections.')])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
for item in [
    '1. Introduction',
    '   1.1 Purpose', '   1.2 Scope', '   1.3 Definitions', '   1.4 References',
    '2. System Architecture Overview',
    '   2.1 Architectural Style', '   2.2 Architecture Diagram',
    '   2.3 Technology Stack', '   2.4 Layer Responsibilities',
    '3. Module Decomposition',
    '   3.1 Component Diagram', '   3.2 Page Modules', '   3.3 UI Components',
    '   3.4 Landing Page Components', '   3.5 Custom Hooks', '   3.6 Service Layer',
    '4. Class & Object Design',
    '   4.1 Class Diagram', '   4.2 Hook Contracts',
    '5. Data Design',
    '   5.1 Firestore Collections', '   5.2 Real-Time Subscription Pattern',
    '6. Interface Design',
    '   6.1 Design System', '   6.2 Navigation & Routing',
    '7. Detailed Component Design (Sequence Diagrams)',
    '   7.1 Authentication & OTP Flow', '   7.2 Create Announcement Flow',
    '   7.3 Interest, NDA & Meeting Flow', '   7.4 Real-Time Chat Flow',
    '   7.5 Admin Audit Flow',
    '8. Security Design',
    '9. Error Handling Strategy',
    '10. Deployment Architecture',
    '11. Build, Configuration & DevOps',
    '12. Performance & Optimisation',
    '13. Testing Strategy',
    '14. Future Work',
]:
    p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)
doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph(
    'This Software Design Document (SDD) describes the internal architecture, module decomposition, '
    'data structures, interaction patterns and deployment design of the HEALTH AI platform. It '
    'translates the requirements specified in the SRS (v3.0) into an implementable design and serves '
    'as the technical blueprint for the development team and any future maintainer. This document '
    'follows the structure recommended by IEEE 1016-2009 Software Design Descriptions.'
)
doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph(
    'The document covers the complete design of the HEALTH AI platform including its React-based '
    'frontend architecture, Firebase Firestore backend integration, real-time communication patterns, '
    'component hierarchy, state management through custom hooks, routing strategy, security mechanisms, '
    'deployment configuration and DevOps pipeline. The design reflects the production build (Version 3) '
    'as of April 2026.'
)
doc.add_heading('1.3 Definitions', level=2)
tbl(['Term', 'Meaning'],
    [('CSR-SPA', 'Client-Side Rendered Single Page Application'),
     ('Hook', 'A React function-based primitive for stateful and side-effectful logic'),
     ('Snapshot', 'A Firestore real-time event delivered through onSnapshot listeners'),
     ('Transaction', 'A Firestore runTransaction call used to atomically read-modify-write multiple documents'),
     ('Subcollection', 'A Firestore-native collection nested under a parent document'),
     ('Liquid-glass', 'The editorial design system used by HEALTH AI — blurred translucent panels with editorial typography')])

doc.add_heading('1.4 References', level=2)
for r in ['HEALTH AI SRS v3.0 (April 2026)',
          'IEEE 1016-2009 Software Design Descriptions',
          'React 19 Documentation — https://react.dev',
          'Firebase Firestore Documentation — https://firebase.google.com/docs/firestore',
          'Framer Motion 12 — https://www.framer.com/motion',
          'Vite 7 — https://vitejs.dev']:
    doc.add_paragraph(r, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. System Architecture Overview', level=1)

doc.add_heading('2.1 Architectural Style', level=2)
doc.add_paragraph(
    'HEALTH AI is implemented as a Client-Side Rendered (CSR) Single Page Application with a '
    'serverless backend (Firebase Firestore). The architecture follows four explicit layers: '
    'Presentation, Business Logic, Data Access and Infrastructure. The following design decisions '
    'are central to the system:'
)
for d in ['Component-Based Architecture: every UI element is a reusable React component.',
          'Hooks-Based State Management: business logic is encapsulated in custom hooks (useAuth, usePosts, useChat, useNotifications, usePostEngagement, useToast) to keep components rendering-focused.',
          'Service Layer Pattern: all Firestore operations are abstracted behind services/firestore.js, providing a single seam for the persistence boundary.',
          'Real-Time First: Firestore onSnapshot listeners drive UI updates instead of polling, enabling sub-second cross-client synchronisation.',
          'Route-Level Code Organisation: each feature owns its own page component under src/pages, lazy-loaded through React Router v7 + Suspense.',
          'Editorial Design System: the liquid-glass.css design system provides visual consistency and supports the prefers-reduced-motion query through Framer Motion MotionConfig.']:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('2.2 Architecture Diagram', level=2)
doc.add_paragraph(
    'The diagram below shows the four-layer architecture of HEALTH AI. Each layer depends only on '
    'the layer immediately beneath it, which keeps the test surface narrow and the upgrade path for '
    'any single technology (e.g., swapping Firestore) localised.'
)
img_diag('architecture_diagram.png', 'Layered Architecture — Presentation, Business Logic, Data Access, Infrastructure')

doc.add_heading('2.3 Technology Stack', level=2)
tbl(['Layer', 'Technology', 'Purpose'],
    [('UI Framework', 'React 19.2', 'Component-based rendering with virtual DOM'),
     ('Build Tool', 'Vite 7.3', 'Fast HMR dev server and optimised production builds'),
     ('Routing', 'React Router v7', 'Declarative client-side routing with guard wrappers'),
     ('Animation', 'Framer Motion 12', 'Physics-based animations, scroll-triggered reveals, MotionConfig'),
     ('3D / WebGL', 'Three.js 0.183 + React Three Fiber + Drei + Spline', 'Interactive 3D landing background, particle effects'),
     ('Video', 'HLS.js', 'Adaptive-bitrate looping video backgrounds with crossfade'),
     ('Icons', 'Lucide React', 'Tree-shakable SVG icon library'),
     ('Database', 'Firebase Firestore v12', 'Real-time NoSQL document database with subcollections'),
     ('Styling', 'Vanilla CSS + CSS Variables + liquid-glass.css', 'Editorial design system with warm amber / cool cyan palette'),
     ('Hashing', 'Web Crypto API', 'SHA-256 password hashing client-side'),
     ('Email', 'EmailJS', 'OTP and meeting notification emails'),
     ('Utilities', 'classnames, tailwind-merge', 'Conditional class composition')])

doc.add_heading('2.4 Layer Responsibilities', level=2)
tbl(['Layer', 'Responsibilities', 'Owns'],
    [('Presentation', 'Renders JSX, dispatches user actions, applies design tokens, animations and accessibility attributes', 'Pages, components, CSS'),
     ('Business Logic', 'Orchestrates state, side effects and validation; owns session, post lifecycle, chat semantics, notifications', 'Custom hooks, context providers'),
     ('Data Access', 'Encapsulates Firestore reads/writes, subscriptions, atomic transactions, presence and audit logging', 'services/firestore.js'),
     ('Infrastructure', 'Provides persistence, real-time channel, email transport and external meeting hosts', 'Firebase Firestore, EmailJS, Zoom/Teams')])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. MODULE DECOMPOSITION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Module Decomposition', level=1)

doc.add_heading('3.1 Component Diagram', level=2)
doc.add_paragraph(
    'The diagram below groups every module of the application into clusters: Pages, Custom Hooks, '
    'Service Layer, Shell / UI Components, Landing Components, External Services, Routing and '
    'Build & Deploy. Arrows indicate the principal "uses" dependencies between clusters.'
)
img_diag('component_diagram.png', 'Component Diagram — clusters and inter-cluster dependencies')

doc.add_heading('3.2 Page Modules', level=2)
tbl(['Module', 'File', 'Responsibility'],
    [('LandingPage', 'pages/LandingPage.jsx', 'Cinematic scrollytelling landing — Hero, BigTextReveal, StickyShowcase, HowItWorks, TwoSides, BentoFeatures, FinalCTA, Footer'),
     ('Login', 'pages/Login.jsx', 'Sign In / Register form with .edu validation, SHA-256 hashing and 6-digit OTP via EmailJS'),
     ('Dashboard', 'pages/Dashboard.jsx', 'Announcement feed with search/filter, 3D tilt cards, bookmarks and match indicators'),
     ('CreatePost', 'pages/CreatePost.jsx', '3-step wizard: Core Details → Technical Info → Settings with WizardProgress'),
     ('PostDetail', 'pages/PostDetail.jsx', 'Two-column detail (content + sidebar). Hosts NDAModal, MeetingSlotsModal and WorkflowActionPanel'),
     ('MyPosts', 'pages/MyPosts.jsx', 'Author post management — stats, filter tabs, status transitions'),
     ('Profile', 'pages/Profile.jsx', 'User card, profile editor, GDPR export & delete controls'),
     ('AdminDashboard', 'pages/AdminDashboard.jsx', '4-tab admin: Overview, Users, Posts, Audit Logs (CSV export)'),
     ('Chat', 'pages/Chat.jsx', 'Two-panel messaging with typing indicators, read receipts, reactions and replies')])

doc.add_heading('3.3 UI Components', level=2)
tbl(['Component', 'Responsibility'],
    [('Navbar', 'Fixed nav bar with logo, links, notifications bell and user menu'),
     ('Notifications', 'Dropdown notification panel with dismiss / dismiss-all actions'),
     ('UserMenu', 'Animated dropdown with profile, shortcuts and logout'),
     ('CommandPalette', 'Global Cmd/Ctrl+K palette for navigation and search'),
     ('ShortcutsModal', 'Keyboard shortcuts reference modal (Shift+?)'),
     ('ToastProvider', 'System-wide toast notifications for critical events'),
     ('LandingBackground', 'Spline-based 3D background, persistent across routes via display:none toggle'),
     ('VideoBackground', 'HLS adaptive-bitrate looping video with crossfade, hidden on landing'),
     ('PageTransition', 'Framer Motion wrapper for page route change animations'),
     ('SkeletonLoader / SkeletonCard', 'Animated placeholder cards during data loading'),
     ('NetworkStatus', 'Offline warning banner indicator'),
     ('GlobalErrorBoundary', 'React Error Boundary catching rendering crashes with recovery UI'),
     ('WizardProgress', 'Animated step indicator for multi-step forms'),
     ('PxSelect', 'Custom styled dropdown select component'),
     ('AnimatedCounter', 'Animated number counter for stats / metrics'),
     ('ScrollToTop', 'Automatic scroll reset on route change'),
     ('ReadingProgress', 'Scroll progress bar for long-form content')])

doc.add_heading('3.4 Landing Page Components', level=2)
tbl(['Component', 'Responsibility'],
    [('Hero', 'Full-viewport value proposition with gradient text and CTA buttons'),
     ('BigTextReveal', 'Per-line clip-path mask text reveal on scroll'),
     ('StickyShowcase', '3-act cinematic pinned scroll experience (160 vh)'),
     ('HowItWorks', '4-step workflow with animated SVG connector path'),
     ('TwoSides', 'Split-pane persona view: Engineer ↔ Clinician'),
     ('BentoFeatures', '4 feature cards in bento grid layout'),
     ('SectionLabel', 'Editorial hairline separators (01 — PREMISE style)'),
     ('SectionNavDots', 'Fixed viewport-aware navigation dots on desktop'),
     ('FinalCTA', 'Call-to-action section before footer'),
     ('Footer', 'Site footer with links and branding'),
     ('ProductMockups', 'SVG mockup illustrations for StickyShowcase')])

doc.add_heading('3.5 Custom Hooks', level=2)
tbl(['Hook', 'State Managed', 'Key Operations'],
    [('useAuth', 'user (from localStorage), session timer', 'login, logout, updateUser, deleteUser, 30-min inactivity timeout'),
     ('usePosts', 'posts, postsLoading', 'addPost, updatePost, updatePostStatus, addInterest, addMeetingRequest, respondToMeeting, auto expiry check'),
     ('useChat', 'messages, conversations, activeRecipient, otherIsTyping', 'send, setTyping, clearHistory, deleteConvo, reactions, replies'),
     ('useNotifications', 'notifications, unreadCount', 'addNotification, dismissNotification, dismissAllNotifications, onNew callback for toasts'),
     ('usePostEngagement', 'interests, meetings (subcollections)', 'Real-time listeners for per-post interests and meetings subcollections'),
     ('useToast', '(context consumer)', 'success, info, error toast dispatchers'),
     ('useAnimReady', 'animReady flag', 'Prevents animations before first paint'),
     ('useInteractiveFX', 'mouse position, hover state', 'Pointer-following glow / tilt effects for cards')])

doc.add_heading('3.6 Service Layer (services/firestore.js)', level=2)
for s in [
    'Posts CRUD: getPosts, addPostToFirestore, updatePostInFirestore, deletePostFromFirestore, subscribeToPostsRT',
    'Post Interests subcollection: addInterestToSubcol, subscribeToPostInterests — uses runTransaction for atomic counter increment',
    'Post Meetings subcollection: addMeetingToSubcol, updateMeetingStatus, subscribeToPostMeetings — uses runTransaction',
    'Users CRUD: getUsers, getUserById, getUserByEmail, emailExists, addUserToFirestore, updateUserInFirestore, deleteUserFromFirestore, subscribeToUsersRT',
    'Activity Logs: addActivityLog, getActivityLogs, subscribeToLogsRT',
    'Notifications: addNotificationToFirestore, deleteNotificationFromFirestore, clearAllNotificationsFromFirestore, subscribeToNotificationsRT',
    'Chat: getOrCreateConversation, sendMessage, deleteMessage, toggleMessageReaction, subscribeToConversationsRT, subscribeToMessagesRT, setTypingStatus, markAsRead, deleteConversation',
    'Presence: updateUserStatus (isOnline + lastSeen serverTimestamp)',
    'Security: hashPassword (SHA-256 via Web Crypto API)',
    'Seeding: seedCollection — populates empty collections with demo data on first run',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. CLASS & OBJECT DESIGN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. Class & Object Design', level=1)

doc.add_heading('4.1 Class Diagram', level=2)
doc.add_paragraph(
    'Although React encourages functional composition, the diagram below uses class-diagram notation '
    'to make the data contracts and dependencies between custom hooks, the firestore.js service and '
    'the persisted entities explicit.'
)
img_diag('class_diagram.png', 'Class Diagram — hooks, service layer and persisted entities')

doc.add_heading('4.2 Hook Contracts', level=2)
tbl(['Hook', 'Returned API', 'Notes'],
    [('useAuth()', '{ user, isAdmin, login, logout, updateUser, deleteUser }', 'Bootstraps from localStorage, starts inactivity timer'),
     ('usePosts()', '{ posts, postsLoading, addPost, updatePost, updatePostStatus, addInterest, addMeetingRequest, respondToMeeting }', 'Subscribes to posts collection in real time'),
     ('useChat(currentUserId)', '{ conversations, messages, activeRecipient, otherIsTyping, send, setTyping, clearHistory, deleteConvo, toggleReaction, reply }', 'Manages conversation list and active conversation'),
     ('useNotifications(userId, onNew?)', '{ notifications, unreadCount, addNotification, dismiss, dismissAll }', 'Subscribes to notifications/{userId} in real time'),
     ('usePostEngagement(postId)', '{ interests, meetings }', 'Subscribes to subcollections with cleanup on unmount'),
     ('useToast()', '{ success, info, error, clear }', 'Context consumer for ToastProvider')])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. DATA DESIGN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. Data Design', level=1)

doc.add_heading('5.1 Firestore Collections', level=2)
tbl(['Collection', 'Document ID', 'Key Fields', 'RT Listener'],
    [('users', 'Custom (admin-TS)', 'name, email, passwordHash, role, institution, city, country, status, isOnline, lastSeen, savedPosts', 'subscribeToUsersRT'),
     ('posts', 'Custom (post-TS)', 'title, domain, projectStage, explanation, confidentiality, status, interestCount, meetingCount, authorId, authorName, authorEmail', 'subscribeToPostsRT'),
     ('posts/{id}/interests', 'Custom (int-UID-TS)', 'userId, userName, message, ndaAccepted, createdAt (serverTimestamp)', 'subscribeToPostInterests'),
     ('posts/{id}/meetings', 'Custom (meet-UID-TS)', 'proposedBy, proposedByName, slot, status, createdAt (serverTimestamp)', 'subscribeToPostMeetings'),
     ('conversations', 'Sorted member IDs joined by _', 'members[], memberData{}, lastMessage, updatedAt, unreadCount{}, isTyping{}', 'subscribeToConversationsRT'),
     ('conversations/{id}/messages', 'Auto-generated', 'senderId, senderName, text, timestamp, read, reactions{}, replyTo{}', 'subscribeToMessagesRT'),
     ('activityLogs', 'Custom (log-TS)', 'timestamp, userId, userName, role, actionType, targetEntity, result, details', 'subscribeToLogsRT'),
     ('notifications', 'Custom ID', 'type, title, message, postId, fromUser, timestamp, read', 'subscribeToNotificationsRT')])

doc.add_heading('5.2 Real-Time Subscription Pattern', level=2)
for i, step in enumerate(['Component mounts → hook initialises → Firestore onSnapshot listener established',
                           'External change occurs → Firestore pushes snapshot to all listeners',
                           'Listener callback fires → hook updates React state via setState',
                           'React re-renders consuming components → UI reflects change instantly',
                           'Component unmounts → unsubscribe() called → listener detached'], 1):
    doc.add_paragraph(f'{i}. {step}')
doc.add_paragraph(
    'Atomic counters (interestCount, meetingCount) are updated through runTransaction to remain '
    'consistent under concurrent writes. Conversations use deterministic IDs derived from the sorted '
    'member IDs to avoid duplicate threads between the same two users.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. INTERFACE DESIGN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. Interface Design', level=1)

doc.add_heading('6.1 Design System', level=2)
doc.add_paragraph(
    'The UI uses a premium editorial liquid-glass design system (liquid-glass.css) with a warm amber '
    'and cyan accent palette, an SVG noise grain overlay and CSS custom properties for theming.'
)
tbl(['Variable', 'Dark Value', 'Purpose'],
    [('--background', '#0a0a0f', 'Main page background'),
     ('--surface', 'rgba(18,18,28,0.6)', 'Card and panel backgrounds'),
     ('--accent-warm', '#f5c48a / #f39a54', 'Primary warm amber accents'),
     ('--accent-cool', '#67e8f9 / #22d3ee', 'Secondary cool cyan accents'),
     ('--text-main', '#e8e8ed', 'Primary text colour'),
     ('--text-muted', 'rgba(255,255,255,0.5)', 'Muted / secondary text'),
     ('--glass-bg', 'rgba(255,255,255,0.04)', 'Liquid-glass panel fill with backdrop-blur'),
     ('fx-grain-global', '3.5% SVG noise', 'Anti-banding texture overlay (Stripe / Arc pattern)')])
img_screen('landing_page.png', 'Landing Page — Editorial liquid-glass design with Spline 3D background')

doc.add_heading('6.2 Navigation & Routing', level=2)
tbl(['Nav Item', 'Route', 'Visibility'],
    [('Feed', '/dashboard', 'All authenticated users'),
     ('My Posts', '/my-posts', 'All authenticated users'),
     ('Messages', '/chat', 'All authenticated users'),
     ('Profile', '/profile', 'All authenticated users'),
     ('Admin Logs', '/admin', 'Admin role only')])
doc.add_paragraph(
    'AppRoutes.jsx wraps protected routes in ProtectedRoute (redirects to /login when no user) and '
    'public routes in GuestRoute (redirects authenticated users away from /login). Pages are loaded '
    'with React.lazy() + Suspense and wrapped in AnimatePresence for transition animations.'
)
img_screen('dashboard_page.png', 'Dashboard — Navbar with active route indicator')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. DETAILED COMPONENT DESIGN — SEQUENCE DIAGRAMS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. Detailed Component Design', level=1)
doc.add_paragraph(
    'This section documents the most important runtime interactions in the application using UML '
    'sequence diagrams. Solid arrows represent synchronous calls; dashed green arrows represent '
    'asynchronous responses or pushed snapshots.'
)

doc.add_heading('7.1 Authentication & OTP Flow', level=2)
doc.add_paragraph(
    'On Sign In, Login.jsx calls useAuth.login() which delegates to firestore.js getUserByEmail(); '
    'the SHA-256 hash is verified, an activity log is written, and the session is established. On '
    'Register, the form submits, an OTP is generated and dispatched through EmailJS, the user '
    'enters the code, and only on success is the user document created in Firestore and the user '
    'signed in.'
)
img_diag('seq_login.png', 'Sequence — Registration with OTP and Sign In')

doc.add_heading('7.2 Create Announcement Flow', level=2)
doc.add_paragraph(
    'CreatePost.jsx renders a 3-step wizard backed by WizardProgress. Each Next click advances the '
    'step state. On final submit, usePosts.addPost() is invoked which calls firestore.js '
    'addPostToFirestore(); the resulting Firestore snapshot is pushed back through '
    'subscribeToPostsRT, the dashboard hook updates state, and React re-renders. The user is then '
    'redirected to /dashboard.'
)
img_diag('seq_create_post.png', 'Sequence — Create Announcement Wizard end-to-end')

doc.add_heading('7.3 Interest, NDA & Meeting Flow', level=2)
doc.add_paragraph(
    'On Express Interest, PostDetail opens NDAModal. The user checks the NDA box and confirms; '
    'usePostEngagement.addInterest() invokes addInterestToSubcol() inside a Firestore transaction '
    'that increments interestCount atomically. A notification is dispatched to the post author. '
    'Subsequently, the user picks a slot from the auto-generated weekday list; the meeting request '
    'is written to posts/{id}/meetings and the author is notified.'
)
img_diag('seq_interest.png', 'Sequence — NDA + Interest + Meeting Proposal')

doc.add_heading('7.4 Real-Time Chat Flow', level=2)
doc.add_paragraph(
    'When a user types, useChat.setTyping(true) updates conversation.isTyping in Firestore; the '
    'snapshot reaches the other browser and animates the typing dots. On Send, sendMessage adds a '
    'document under conversations/{id}/messages; the snapshot reaches both browsers. When the '
    'recipient opens the conversation, markAsRead() updates message.read; the originating browser '
    'sees a double-check.'
)
img_diag('seq_chat.png', 'Sequence — Real-Time Chat with typing indicator and read receipts')

doc.add_heading('7.5 Admin Audit Flow', level=2)
doc.add_paragraph(
    'AdminDashboard renders four tabs. On Audit Logs, subscribeToLogsRT delivers a real-time view '
    'of activityLogs. The admin filters by role, action type or text. Export CSV iterates the '
    'filtered list, builds a CSV string and triggers a Blob download in the browser. Freeze / '
    'Unfreeze writes to user.status; the next login attempt by a frozen user is rejected.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. SECURITY DESIGN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. Security Design', level=1)
tbl(['Layer', 'Mechanism', 'Detail'],
    [('Password', 'SHA-256 hashing', 'Client-side via Web Crypto API; no plaintext stored or transmitted'),
     ('Email validation', '.edu enforcement', 'Personal providers (gmail, yahoo, outlook, hotmail, icloud, proton, yandex, mail.ru, etc.) blocked'),
     ('Email verification', '6-digit OTP via EmailJS', '10-minute TTL, 60-second resend cooldown; account inactive until verified'),
     ('Routes', 'Guard components', 'ProtectedRoute redirects unauthenticated users; GuestRoute redirects authenticated users away from /login'),
     ('IP protection', 'NDA workflow', 'Confidential posts gate the technical blueprint behind an explicit NDA acknowledgement modal'),
     ('RBAC', 'Role-based access', 'Cross-role discovery only; admin routes restricted to Admin role'),
     ('Audit', 'Activity logging', 'All login, post, meeting and admin actions persisted to activityLogs with timestamp, user and result'),
     ('GDPR', 'Data rights', 'Article 15 (access via Profile), Article 17 (account delete), Article 20 (JSON export)'),
     ('Session', 'Inactivity timeout', '30-minute timeout on user inactivity; cleared from localStorage'),
     ('Transport', 'TLS', 'All network traffic over HTTPS; Firestore SDK uses TLS WebSocket')])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9. ERROR HANDLING STRATEGY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('9. Error Handling Strategy', level=1)
tbl(['Error Type', 'Detection', 'User Experience'],
    [('Render crash', 'GlobalErrorBoundary (componentDidCatch)', 'Recovery UI with reload button instead of blank screen'),
     ('Network loss', 'NetworkStatus + navigator.onLine + window listeners', 'Offline warning banner; app remains responsive for cached views'),
     ('Auth failure', 'Try/catch in handlers + activity log', 'Animated inline error message; non-leaking error copy'),
     ('Firestore error', 'Try/catch around all service calls + console.error', 'Toast notification with retry hint; graceful degradation of the affected view'),
     ('Form validation', 'Client-side validation before submit', 'Contextual inline error messages adjacent to the offending field'),
     ('OTP errors', 'EmailJS response code + TTL check', 'Inline message explaining expired or invalid code; resend cooldown enforced')])

# ══════════════════════════════════════════════════════════════════
# 10. DEPLOYMENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('10. Deployment Architecture', level=1)
doc.add_paragraph(
    'The HEALTH AI platform is deployed as a static SPA bundle served from a CDN; all backend '
    'concerns are handled by Google Firebase. The diagram below shows the runtime nodes and the '
    'protocols connecting them.'
)
img_diag('deployment_diagram.png', 'Deployment Diagram — Browser, Static Host, Firebase Cloud and external services')

doc.add_paragraph(
    'Build pipeline: source → Vite production build → dist/ (HTML, CSS, JS bundles) → static host. '
    'Runtime: browser ↔ Firestore SDK (WebSocket) ↔ Firebase Firestore (Google Cloud).'
)

# ══════════════════════════════════════════════════════════════════
# 11. BUILD, CONFIGURATION & DEVOPS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('11. Build, Configuration & DevOps', level=1)
tbl(['Concern', 'How it is handled'],
    [('Local dev', '`npm run dev` starts Vite at http://localhost:5173 with HMR'),
     ('Production build', '`npm run build` emits a tree-shaken, code-split static bundle to dist/'),
     ('Preview', '`npm run preview` serves the production bundle locally for smoke testing'),
     ('Configuration', '.env (untracked) contains VITE_FIREBASE_API_KEY, VITE_FIREBASE_AUTH_DOMAIN, VITE_FIREBASE_PROJECT_ID, VITE_FIREBASE_STORAGE_BUCKET, VITE_FIREBASE_MESSAGING_SENDER_ID, VITE_FIREBASE_APP_ID; example values in README'),
     ('Routing fallback', 'SPA fallback rewrites must point all unknown paths to index.html on the host'),
     ('Static hosting', 'Compatible with Firebase Hosting, Vercel or Netlify; no server-side runtime needed'),
     ('Screenshots', 'take-screenshots.js (Playwright) refreshes docs/screenshots/ from the running dev server'),
     ('Documentation generation', 'docs/generate_srs.py, generate_sdd.py and generate_userguide.py emit the .docx artefacts and embed live screenshots and matplotlib-generated diagrams')])

# ══════════════════════════════════════════════════════════════════
# 12. PERFORMANCE & OPTIMISATION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('12. Performance & Optimisation', level=1)
for p in ['Code splitting through React.lazy() at the page level keeps the initial bundle minimal.',
          'Vite\'s production build performs tree shaking and dynamic import chunking automatically.',
          'Firestore subscriptions reuse a single WebSocket per app instance, multiplexing queries.',
          'Counters are denormalised on the parent document (interestCount, meetingCount) to avoid expensive aggregate queries.',
          'Conversations use deterministic IDs derived from sorted member IDs; the SDK can resolve them in a single document read.',
          'Animations use the GPU-accelerated transform/opacity properties; layout thrash is avoided.',
          'HLS video backgrounds load only on routes that need them and crossfade with low-CPU transitions.',
          'The Spline 3D background is mounted once and toggled with display:none on non-landing routes to preserve the WebGL context.',
          'prefers-reduced-motion is honoured globally through MotionConfig to keep CPU usage low for sensitive users.']:
    doc.add_paragraph(p, style='List Bullet')

# ══════════════════════════════════════════════════════════════════
# 13. TESTING STRATEGY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('13. Testing Strategy', level=1)
tbl(['Test Layer', 'Approach', 'Coverage Focus'],
    [('Manual smoke', 'Two-account walkthrough on Chrome and Firefox', 'End-to-end happy paths for UC-01..UC-08'),
     ('Cross-browser', 'Latest two majors of Chrome, Firefox, Safari, Edge', 'Layout, animation, WebSocket'),
     ('Responsive', 'DevTools at 320, 768, 1024, 1440, 2560 px', 'Navbar collapse, card grid, modals'),
     ('Accessibility spot-check', 'axe-core / Lighthouse', 'Contrast, semantic landmarks, focus order'),
     ('Performance spot-check', 'Lighthouse, devtools throttling', 'Time-to-Interactive < 2 s on Fast 3G'),
     ('Security spot-check', 'OWASP Top 10 review', 'XSS, injection surface, transport security')])

# ══════════════════════════════════════════════════════════════════
# 14. FUTURE WORK
# ══════════════════════════════════════════════════════════════════
doc.add_heading('14. Future Work', level=1)
for f in ['Migrate authentication to Firebase Auth + custom claims to remove client-side hashing entirely.',
          'Introduce server-side Firestore Security Rules tied to the auth uid and role claim.',
          'Add localisation (i18n) — extract UI strings into JSON resource files (English / Turkish to start).',
          'Implement an institution directory and verified institution badges.',
          'Provide a moderation queue with reason codes instead of direct deletion.',
          'Add multi-factor authentication beyond OTP (TOTP / WebAuthn).',
          'Surface admin analytics (cohort, retention, time-to-first-meeting) through a charting library.',
          'Generate ICS calendar invites for accepted meetings and integrate with the user\'s calendar.']:
    doc.add_paragraph(f, style='List Bullet')

# ═══════════════ SAVE ═══════════════
output = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'HEALTH_AI_SDD_Document.docx')
doc.save(output)
print(f'SDD saved to: {output}')
