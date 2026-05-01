#!/usr/bin/env python3
"""Generate HEALTH AI User Guide — comprehensive final version with diagrams and tutorials."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

for lv in range(1, 4):
    h = doc.styles[f'Heading {lv}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h.font.bold = True
    h.font.size = Pt([0, 22, 16, 13][lv])
    if lv == 1:
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif lv == 2:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

SDIR = os.path.join(os.path.dirname(__file__), 'screenshots')
DDIR = os.path.join(os.path.dirname(__file__), 'diagrams')


def img_screen(fn, cap, w=Inches(5.6)):
    fp = os.path.join(SDIR, fn)
    if os.path.exists(fp):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(fp, width=w)
        c = doc.add_paragraph(f'Figure: {cap}')
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.italic = True
        c.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()


def img_diag(fn, cap, w=Inches(6.2)):
    fp = os.path.join(DDIR, fn)
    if os.path.exists(fp):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(fp, width=w)
        c = doc.add_paragraph(f'Diagram: {cap}')
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.italic = True
        c.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()


def tbl(hdrs, rows):
    t = doc.add_table(rows=1, cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(hdrs):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    for rd in rows:
        r = t.add_row()
        for i, tx in enumerate(rd):
            r.cells[i].text = str(tx)
            r.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


def callout(title, body, kind='Tip'):
    p = doc.add_paragraph()
    run = p.add_run(f'{kind} — {title}: ')
    run.bold = True
    if kind == 'Warning':
        run.font.color.rgb = RGBColor(0xb9, 0x1c, 0x1c)
    elif kind == 'Tip':
        run.font.color.rgb = RGBColor(0x07, 0x66, 0x4e)
    else:
        run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p.add_run(body)


# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('HEALTH AI'); r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = RGBColor(0x3b, 0x3b, 0x98)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('European HealthTech Co-Creation & Innovation Platform').font.size = Pt(14)
doc.add_paragraph()
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run('User Guide'); r.font.size = Pt(20); r.font.bold = True
doc.add_paragraph()
v = doc.add_paragraph(); v.alignment = WD_ALIGN_PARAGRAPH.CENTER
v.add_run('Version 3.1 — Final Release • May 2026').font.size = Pt(12)
doc.add_paragraph()

gp = doc.add_paragraph(); gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
gr = gp.add_run('Group: Health Shield'); gr.font.size = Pt(16); gr.font.bold = True; gr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for m in ['Cem Özal', 'Emre Kurubaş', 'Hasabu Can Eltayeb', 'Sertaç Ataç']:
    mr = mp.add_run(m + '\n'); mr.font.size = Pt(12); mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

it = doc.add_table(rows=0, cols=2); it.style = 'Light Grid Accent 1'
for lb, vl in [('Document', 'User Guide'),
               ('Project', 'HEALTH AI Platform'),
               ('Group', 'Health Shield'),
               ('Members', 'Cem Özal, Emre Kurubaş, Hasabu Can Eltayeb, Sertaç Ataç'),
               ('Version', '3.1 (Final)'),
               ('Date', 'May 1, 2026'),
               ('Course', 'SENG 384 — Software Engineering'),
               ('Audience', 'Healthcare Professionals, Engineers, Admins'),
               ('Status', 'Final — comprehensive with diagrams and tutorials')]:
    rw = it.add_row(); rw.cells[0].text = lb; rw.cells[1].text = vl
    rw.cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
for item in ['1. Welcome to HEALTH AI',
             '   1.1 What is HEALTH AI?', '   1.2 Who Should Use This Guide?',
             '   1.3 The Big Picture (Journey Diagram)', '   1.4 Accessing the Platform',
             '2. Creating Your Account',
             '   2.1 Registration Requirements', '   2.2 Step-by-Step Registration',
             '   2.3 OTP Email Verification', '   2.4 Signing In', '   2.5 Forgot Password',
             '3. Navigating the Platform',
             '   3.1 Navigation Bar', '   3.2 Navigation Map (Diagram)',
             '   3.3 Keyboard Shortcuts', '   3.4 Command Palette (Cmd+K)',
             '4. Discovering Projects (Dashboard)',
             '   4.1 Browsing the Feed', '   4.2 Searching and Filtering',
             '   4.3 Saving Projects (Bookmarks)', '   4.4 Match Indicators',
             '5. Creating an Announcement',
             '   5.1 Wizard Overview', '   5.2 Step 1 — Core Details',
             '   5.3 Step 2 — Technical Info', '   5.4 Step 3 — Settings',
             '   5.5 Confidentiality Decision (Diagram)',
             '6. Viewing Announcement Details',
             '   6.1 Layout', '   6.2 Expressing Interest (NDA)', '   6.3 Proposing a Meeting',
             '7. Managing Your Announcements',
             '   7.1 My Posts Overview', '   7.2 Status Transitions',
             '8. Messaging',
             '   8.1 Starting a Conversation', '   8.2 Chat Features',
             '   8.3 Reactions, Replies & Read Receipts',
             '9. Profile Management',
             '   9.1 Editing Your Profile', '   9.2 GDPR Data Rights',
             '   9.3 Privacy Policy & Terms of Service',
             '10. Notifications',
             '11. Admin Features',
             '12. Accessibility & Reduced Motion',
             '13. Troubleshooting & FAQ',
             '14. Contact & Support']:
    p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. WELCOME
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. Welcome to HEALTH AI', level=1)

doc.add_heading('1.1 What is HEALTH AI?', level=2)
doc.add_paragraph(
    'HEALTH AI is a European HealthTech Co-Creation and Innovation Platform that connects healthcare '
    'professionals (doctors, clinicians, researchers) with engineers and developers. The platform '
    'enables structured partner discovery for health-tech innovation projects through a secure, '
    'GDPR-compliant workflow.'
)
doc.add_paragraph('Key benefits:')
for b in ['Find the right partner for your health-tech project across European institutions',
          'Secure NDA-protected communication workflow before any technical detail is shared',
          'GDPR-compliant data handling with full data export and account-deletion rights',
          'Cross-disciplinary matching between healthcare and engineering disciplines',
          'No intellectual property is stored on the platform — it facilitates contact only',
          'Real-time messaging and notifications keep you in the loop without manual refresh',
          'Designed for both clinicians (no engineering background required) and engineers (no clinical background required)']:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('1.2 Who Should Use This Guide?', level=2)
tbl(['Role', 'Description', 'Access'],
    [('Healthcare Professional', 'Doctors, clinicians, researchers with clinical expertise and innovative ideas', 'Full platform: create posts, browse, express interest, chat, profile'),
     ('Engineer / Developer', 'Software, hardware, biomedical engineers with technical capabilities', 'Full platform: create posts, browse, express interest, chat, profile'),
     ('Admin', 'Platform administrator', 'Full access plus user management, post moderation and audit logs')])

doc.add_heading('1.3 The Big Picture (Journey Diagram)', level=2)
doc.add_paragraph(
    'The diagram below shows the typical end-to-end journey from registration to a confirmed '
    'partnership. Each numbered step corresponds to a section of this guide.'
)
img_diag('user_workflow.png', 'User Journey — From Registration to Partnership')

doc.add_heading('1.4 Accessing the Platform', level=2)
doc.add_paragraph(
    'Open your web browser (Chrome, Firefox, Safari or Edge — latest two major versions) and '
    'navigate to the HEALTH AI URL provided by your institution or instructor. You will be greeted '
    'by the cinematic landing page.'
)
img_screen('landing_page.png', 'HEALTH AI Landing Page')
doc.add_paragraph('Click "Join the Network" or "Access Platform" to create an account or sign in.')
callout('Browser support', 'JavaScript and WebSockets must be enabled. WebGL is recommended for the 3D landing background — the page falls back to a static image where WebGL is unavailable.', 'Note')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. CREATING YOUR ACCOUNT
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. Creating Your Account', level=1)

doc.add_heading('2.1 Registration Requirements', level=2)
doc.add_paragraph('To register on HEALTH AI you need:')
for r in ['An institutional email address ending with .edu (Gmail, Yahoo, Hotmail and other personal providers are not accepted)',
          'A password with at least 6 characters',
          'Your full name with professional title (e.g., "Dr.", "Prof.")',
          'Your role: Healthcare Professional or Engineer / Developer',
          'Your institution name (e.g., "MIT", "University of Zurich")',
          'Your country and city',
          'Access to your email inbox to receive the 6-digit OTP verification code']:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('2.2 Step-by-Step Registration', level=2)
steps = [
    'On the landing page, click "Access Platform" / "Join the Network".',
    'On the Login page, click the "Register" tab at the top of the form.',
    'Enter your institutional email address (must end with .edu).',
    'Enter a secure password (minimum 6 characters).',
    'Enter your full name with title (e.g., "Dr. Maria Rossi").',
    'Select your role: "Healthcare Professional" or "Engineer / Developer".',
    'Enter your institution name, country and city.',
    'Click "Create Account". You will be taken to the OTP verification panel.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')
img_screen('login_page.png', 'Login / Registration Page')

doc.add_heading('2.3 OTP Email Verification', level=2)
doc.add_paragraph(
    'After you submit the registration form, the system sends a 6-digit code to the email address '
    'you provided. Open your inbox and copy the code into the OTP panel. The code is valid for '
    '10 minutes; if it expires, click "Resend Code" (a 60-second cooldown applies between resends).'
)
callout('Check your spam folder', 'OTP emails sometimes land in spam or junk. If you cannot find the email after a minute, check there before resending.', 'Tip')
callout('Account is inactive until OTP is verified', 'You cannot sign in with your new credentials until you successfully verify the OTP. The account document is created in our database only on successful verification.', 'Warning')

doc.add_heading('2.4 Signing In', level=2)
doc.add_paragraph('If you already have a verified account:')
for s in ['Open the Login page.',
          'Make sure the "Sign In" tab is selected.',
          'Enter your institutional email and password.',
          'Click "Sign In". You will be redirected to the Dashboard.']:
    doc.add_paragraph(s, style='List Bullet')
callout('Failed attempts are logged', 'Every successful and failed login attempt is recorded in the system audit log so administrators can detect abuse. There is no rate limiter for end users today, but persistent failures can lead to account freeze by an admin.', 'Note')

doc.add_heading('2.5 Forgot Password', level=2)
doc.add_paragraph(
    'A self-service password reset flow is not yet available. If you have forgotten your password, '
    'contact a platform administrator who can either reset it for you or freeze and recreate the '
    'account. Future releases will add an automated reset over OTP.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. NAVIGATING THE PLATFORM
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Navigating the Platform', level=1)

doc.add_heading('3.1 Navigation Bar', level=2)
doc.add_paragraph('Once signed in you will see the navigation bar at the top of every page. It contains:')
tbl(['Element', 'Description', 'Action'],
    [('HEALTH AI Logo', 'Platform branding with gradient icon', 'Click to go to the Dashboard'),
     ('Feed', 'Main announcement dashboard', 'Browse all project announcements'),
     ('My Posts', 'Your posted announcements', 'Manage your own announcements'),
     ('Messages', 'Real-time chat system', 'Send and receive messages'),
     ('Profile', 'Your user profile', 'Edit profile and manage data'),
     ('Admin Logs', 'Admin dashboard (Admin only)', 'System oversight and management'),
     ('Bell icon', 'Notifications', 'View interest and meeting notifications'),
     ('User badge', 'Your name and role', 'Shows your identity and role'),
     ('Logout', 'Sign out', 'Securely end your session')])

doc.add_heading('3.2 Navigation Map (Diagram)', level=2)
doc.add_paragraph(
    'The map below summarises which page hosts which feature. Use it as a quick reference when you '
    'are looking for a specific action.'
)
img_diag('navigation_tree.png', 'Navigation Map — every feature and where to find it')

doc.add_heading('3.3 Keyboard Shortcuts', level=2)
tbl(['Shortcut', 'Action'],
    [('Ctrl + K  /  Cmd + K', 'Open the Command Palette (global navigation and search)'),
     ('Shift + ?', 'Open the keyboard shortcuts reference modal'),
     ('Esc', 'Close any open modal or dropdown'),
     ('Tab / Shift+Tab', 'Move focus between interactive elements'),
     ('Enter', 'Send a message or submit the focused form')])

doc.add_heading('3.4 Command Palette (Cmd + K)', level=2)
doc.add_paragraph(
    'Press Cmd+K (Mac) or Ctrl+K (Windows / Linux) from any authenticated page to open the Command '
    'Palette. Start typing to navigate to any page or to search across announcements. Use the up / '
    'down arrows to highlight a result and Enter to activate it.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. DISCOVERING PROJECTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. Discovering Projects (Dashboard)', level=1)

doc.add_heading('4.1 Browsing the Feed', level=2)
doc.add_paragraph(
    'The Dashboard is your main hub for discovering projects. Here you will find announcements from '
    'users with complementary roles — if you are an Engineer you will see Healthcare Professional '
    'posts, and vice versa.'
)
img_screen('dashboard_page.png', 'Dashboard — Innovator Feed with project cards')
doc.add_paragraph('Each project card displays:')
for c in ['Role badge (Engineer or Healthcare Professional) indicating the author role',
          'Domain badge (e.g., Cardiology, Neurology, Oncology)',
          'Status badge (Active, Meeting Scheduled, Partner Found)',
          'Project title and description excerpt',
          'Match indicators (same city, complementary role)',
          'Bookmark button to save the project for later']:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('4.2 Searching and Filtering', level=2)
doc.add_paragraph('Use the search bar and filter dropdowns to narrow down the feed:')
for f in ['Search: type keywords to search across titles, descriptions, domains and expertise fields',
          'Domain: filter by medical or technical domain',
          'Stage: filter by project stage (Idea, Concept Validation, Prototype, Pilot, Pre-Deployment)',
          'Country: filter by the project country',
          'City: filter by the project city',
          'Status: filter by announcement status (default: Active only)']:
    doc.add_paragraph(f, style='List Bullet')
doc.add_paragraph('A counter in the top-right of the filter bar shows how many announcements match your current criteria.')

doc.add_heading('4.3 Saving Projects (Bookmarks)', level=2)
doc.add_paragraph(
    'Click the bookmark icon on any project card to save it. Switch to the "Saved Projects" tab to '
    'view all of your bookmarked projects. Bookmarks are personal to your account and persist '
    'across devices.'
)

doc.add_heading('4.4 Match Indicators', level=2)
doc.add_paragraph(
    'When the dashboard detects useful overlap between you and a post (same city, complementary '
    'role) it shows a match indicator badge on the card. Use these as quick signals for high-fit '
    'projects, but always read the full post detail before expressing interest.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. CREATING AN ANNOUNCEMENT
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. Creating an Announcement', level=1)

doc.add_heading('5.1 Wizard Overview', level=2)
doc.add_paragraph(
    'Click the "+ New Announcement" button on the Dashboard to start creating an announcement. The '
    'form is structured as a 3-step wizard with a visual progress indicator. You can use the Back / '
    'Next buttons to move between steps; nothing is saved until you click "Publish Announcement" on '
    'the final step.'
)

doc.add_heading('5.2 Step 1 — Core Details', level=2)
for f in ['Title: give your project a clear, descriptive title (e.g., "AI-Assisted ECG Analysis for Arrhythmia Detection")',
          'Domain: select the relevant medical or technical domain from the dropdown',
          'Project Stage: choose your current stage by clicking one of the cards (Idea, Concept Validation, Prototype Developed, Pilot Testing, Pre-Deployment)',
          'Description: describe the problem you are solving, initial outcomes and the ideal partner profile']:
    doc.add_paragraph(f, style='List Bullet')
img_screen('create_post_page.png', 'Create Announcement — Step 1: Core Details')

doc.add_heading('5.3 Step 2 — Technical Info', level=2)
for f in ['Technical Blueprint (optional): describe your technical approach. If you set the post as NDA-protected this section will only be visible to viewers who accept the NDA.',
          'Required Expertise: describe what kind of partner you are looking for.',
          'Collaboration Type: select Co-Development, Advisory, Licensing, Joint Research or Pilot Partnership.']:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.4 Step 3 — Settings', level=2)
for f in ['Country / City: set the project location.',
          'Confidentiality Level: choose "Public Info" (visible to all) or "NDA Protected" (technical blueprint hidden until NDA acceptance).',
          'Expiry: choose how long the announcement remains active (15, 30, 60 or 90 days).']:
    doc.add_paragraph(f, style='List Bullet')
doc.add_paragraph('Click "Publish Announcement" to make your post live on the platform.')
callout('You can edit the title and description later', 'Open the post detail page from My Posts; an Edit dialog lets you adjust the title and description after publish.', 'Tip')

doc.add_heading('5.5 Confidentiality Decision (Diagram)', level=2)
doc.add_paragraph(
    'When choosing between Public Info and NDA Protected, the diagram below summarises how each '
    'choice affects what other users see on the post detail page.'
)
img_diag('nda_decision_flow.png', 'NDA Visibility Decision Flow')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. VIEWING ANNOUNCEMENT DETAILS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. Viewing Announcement Details', level=1)
doc.add_paragraph('Click on any project card from the Dashboard to view its full detail page.')

doc.add_heading('6.1 Layout', level=2)
doc.add_paragraph('The detail page is organised in two columns:')
for i in ['Status badges, domain and author role at the top',
          'Project metadata: stage, location, posting date and expiry date',
          'Executive Overview: full project description',
          'Technical Blueprint: detailed technical approach (may be NDA-gated)',
          'Required Expertise: what the author is looking for in a partner',
          'Sidebar: target counterparty info, collaboration type and data sharing level',
          'Author card with a "Message" button for direct chat']:
    doc.add_paragraph(i, style='List Bullet')
img_screen('post_detail_page.png', 'Post Detail — Two-column layout with sidebar')

doc.add_heading('6.2 Expressing Interest (NDA)', level=2)
doc.add_paragraph('To express interest in a project:')
steps = ['Click the "Express Interest" button in the sidebar.',
         'An NDA acceptance modal will appear. Read the non-disclosure terms carefully.',
         'Tick the agreement checkbox to accept the NDA terms.',
         'Optionally write a short message to the post author.',
         'Click "Confirm Interest". The post author will receive a notification.']
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')
callout('NDA-protected posts', 'For NDA-protected posts the Technical Blueprint section becomes visible only after you confirm interest and accept the NDA.', 'Note')

doc.add_heading('6.3 Proposing a Meeting', level=2)
doc.add_paragraph('After expressing interest you can propose a meeting:')
for s in ['Click "Propose Meeting" on the post detail page.',
          'Select one of the auto-generated time slots (weekday options for the next 5 business days at 10:00, 14:00 and 16:00).',
          'The post author receives a notification and can Accept or Decline the request.',
          'On accept, the post status changes to "Meeting Scheduled" with instructions to use Zoom or Microsoft Teams for the actual meeting.']:
    doc.add_paragraph(s, style='List Bullet')
callout('No video calls inside HEALTH AI', 'The platform does not host video calls. Use Zoom or Microsoft Teams for the actual meeting; HEALTH AI only facilitates the scheduling.', 'Warning')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. MANAGING YOUR ANNOUNCEMENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. Managing Your Announcements', level=1)

doc.add_heading('7.1 My Posts Overview', level=2)
doc.add_paragraph('Go to "My Posts" in the navigation bar to manage your announcements. Available features:')
for f in ['Statistics dashboard: total, drafts, active, in-meeting and closed counts at a glance',
          'Filter tabs: filter your posts by status (All, Draft, Active, Meeting Scheduled, Partner Found)',
          'Publish: publish drafts to make them visible on the platform',
          'Close: mark a post as Partner Found when you have found a partner',
          'View: open the full post detail page']:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('7.2 Status Transitions', level=2)
doc.add_paragraph('A post moves through the following states during its lifecycle:')
tbl(['State', 'Meaning', 'How to reach it'],
    [('Active', 'Live and visible on the dashboard', 'Default after Publish'),
     ('Active + Interest Received', 'At least one viewer has expressed interest', 'Automatically once any user submits interest'),
     ('Meeting Scheduled', 'You have accepted a meeting request', 'Click Accept on a meeting request'),
     ('CLOSED (Partner Found)', 'You have found a partner and closed the post', 'Click "Mark as Partner Found"'),
     ('Expired', 'Past the configured expiry date', 'Automatic'),
     ('DELETED', 'Removed by the admin or owner', 'Admin moderation or owner delete')])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. MESSAGING
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. Messaging', level=1)
doc.add_paragraph(
    'HEALTH AI includes a built-in real-time messaging system for secure communication between users.'
)

doc.add_heading('8.1 Starting a Conversation', level=2)
for w in ['From a post: click the "Message" button on the author card in the post detail page.',
          'From the Messages page: click the chat icon in the sidebar header, then select a user from the discovery list.']:
    doc.add_paragraph(w, style='List Bullet')
img_screen('chat_page.png', 'Messages Page — Chat interface')

doc.add_heading('8.2 Chat Features', level=2)
for f in ['Real-time messages: messages appear instantly for both participants',
          'Typing indicators: animated dots when the other person is typing',
          'Read receipts: single check (sent) and double check (read)',
          'Date groups: messages are organised by date for easy scanning',
          'Search: filter your conversations using the sidebar search bar',
          'Chat management: three-dot menu offers Clear History and Delete Conversation',
          'Unread badges: unread message count is shown next to each conversation']:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('8.3 Reactions, Replies & Read Receipts', level=2)
doc.add_paragraph(
    'Hover over a message to access an emoji reaction picker; click a reaction to toggle it. Use '
    'the reply icon next to a message to quote it in your next message — the original snippet is '
    'shown above your reply for context. Read receipts update the moment the recipient opens the '
    'conversation.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9. PROFILE MANAGEMENT
# ══════════════════════════════════════════════════════════════════
doc.add_heading('9. Profile Management', level=1)
doc.add_paragraph('Go to "Profile" in the navigation bar to view and manage your account.')
img_screen('profile_page.png', 'Profile Page — User info, settings and GDPR controls')

doc.add_heading('9.1 Editing Your Profile', level=2)
for s in ['Click the "Edit" button in the Profile Settings section.',
          'Update your Full Name, Institution, Country or City.',
          'Click "Save" to apply changes, or "Cancel" to discard.',
          'A success message will confirm your profile has been updated.']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('9.2 GDPR Data Rights', level=2)
doc.add_paragraph('HEALTH AI is fully GDPR-compliant. You have the following rights:')
tbl(['Right', 'GDPR Article', 'How to use'],
    [('Right of Access', 'Article 15', 'Your profile page shows all personal data the system holds about you'),
     ('Right to Data Portability', 'Article 20', 'Click "Export My Data (JSON)" to download all your data as a JSON file'),
     ('Right to Erasure', 'Article 17', 'Click "Delete My Account (Art. 17)" to permanently delete your account and all associated personal data')])
callout('Account deletion is permanent', 'Deletion cannot be undone. Your posts will become anonymised or removed depending on platform policy. Export your data first if you want a personal archive.', 'Warning')

doc.add_heading('9.3 Privacy Policy & Terms of Service', level=2)
doc.add_paragraph(
    'Two public legal pages are accessible without signing in via the footer of every page:'
)
tbl(['Page', 'Route', 'Contents'],
    [('Privacy Policy', '/privacy', 'How HEALTH AI collects, processes, stores and deletes personal data; lawful basis under GDPR; your rights under Articles 15, 17 and 20'),
     ('Terms of Service', '/terms', 'Acceptable-use rules, account responsibilities, NDA acknowledgement scope, intellectual-property handling and limitation of liability')])
doc.add_paragraph(
    'During registration you implicitly accept the latest version of both documents. Material '
    'changes are communicated by an in-app banner the next time you sign in.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 10. NOTIFICATIONS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('10. Notifications', level=1)
doc.add_paragraph(
    'The bell icon in the navigation bar shows a count badge whenever you have unread '
    'notifications. Click it to expand the dropdown and review:'
)
for n in ['New interest expressions on your posts',
          'New meeting requests from interested users',
          'Responses to your meeting requests (Accept / Decline)',
          'System notifications about your account or posts']:
    doc.add_paragraph(n, style='List Bullet')
doc.add_paragraph(
    'Click an individual notification to dismiss it, or click "Dismiss All" at the bottom of the '
    'dropdown to clear them in one action. Critical events also surface as a toast in the bottom '
    'corner of the screen.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 11. ADMIN FEATURES
# ══════════════════════════════════════════════════════════════════
doc.add_heading('11. Admin Features', level=1)
doc.add_paragraph(
    'If you have the Admin role you can access the Admin Dashboard via "Admin Logs" in the '
    'navigation bar. The admin panel provides four tabs:'
)
tbl(['Tab', 'Features'],
    [('Overview', 'Dashboard cards showing total users, total announcements, active projects and monitored events'),
     ('Users', 'User management table with name, email, role, status (Active / Frozen) and freeze / unfreeze action'),
     ('Announcements', 'Post moderation table with title, author, status, privacy level and delete action'),
     ('Audit Logs', 'Searchable and filterable audit log table with timestamp, user, action type and details. Supports CSV export for compliance reporting.')])

doc.add_paragraph('Admin capabilities:')
for c in ['Freeze / Unfreeze users to temporarily disable or re-enable an account',
          'Delete posts to remove any announcement from the platform',
          'Export audit logs as CSV for compliance and incident response',
          'Monitor activity: logins, post creation, meeting requests and failed attempts']:
    doc.add_paragraph(c, style='List Bullet')
callout('Becoming an admin', 'The first account that registers with admin@healthai.edu is automatically promoted to Admin. Additional admins must be granted by an existing admin.', 'Note')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 12. ACCESSIBILITY & REDUCED MOTION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('12. Accessibility & Reduced Motion', level=1)
doc.add_paragraph(
    'HEALTH AI is built with accessibility in mind. Interactive elements are reachable by keyboard, '
    'colour contrast targets WCAG AA, and the layout adapts from mobile to wide desktop viewports.'
)
doc.add_paragraph(
    'Animations honour the operating-system "reduce motion" preference. To enable it:'
)
tbl(['Operating System', 'How to enable'],
    [('macOS', 'System Settings → Accessibility → Display → Reduce motion'),
     ('Windows 10 / 11', 'Settings → Accessibility → Visual effects → Animation effects: Off'),
     ('iOS / iPadOS', 'Settings → Accessibility → Motion → Reduce Motion'),
     ('Android', 'Settings → Accessibility → Visibility enhancements → Remove animations')])
doc.add_paragraph(
    'When reduce-motion is on, scroll-triggered reveals, page transitions and the BigTextReveal '
    'animations are skipped. Functional behaviour is identical.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 13. TROUBLESHOOTING & FAQ
# ══════════════════════════════════════════════════════════════════
doc.add_heading('13. Troubleshooting & FAQ', level=1)
tbl(['Question', 'Answer'],
    [('I cannot register with my email', 'Only institutional .edu addresses are accepted. Personal providers (Gmail, Yahoo, Hotmail, Outlook, iCloud, ProtonMail, etc.) are blocked.'),
     ('My OTP code does not arrive', 'Wait up to a minute and check your spam folder. If still missing, click "Resend Code" (60-second cooldown applies).'),
     ('My OTP code says "expired"', 'Codes are valid for 10 minutes. Click "Resend Code" to get a new one.'),
     ('I forgot my password', 'A self-service reset is not yet available. Contact a platform administrator.'),
     ('I cannot see the Technical Blueprint', 'The post is NDA-protected. Express interest and accept the NDA to unlock it.'),
     ('My post is not showing on the Dashboard', 'Check the status filter — the default is "Active" only. Your post may be in Draft.'),
     ('I am offline and features are not working', 'HEALTH AI requires an internet connection for real-time features. The offline indicator banner appears when connectivity is lost.'),
     ('The 3D landing background looks broken', 'Your browser may have WebGL disabled. The page falls back to a static image; functionality is unaffected.'),
     ('My typing indicator is stuck', 'Refresh the page; the typing flag clears automatically when the conversation closes.'),
     ('How do I delete my account?', 'Profile → Data Privacy (GDPR) → "Delete My Account (Art. 17)". This action is permanent.'),
     ('Can I change my role?', 'Role changes are not supported through the UI. Contact an administrator.'),
     ('Where do meetings take place?', 'Externally, on Zoom or Microsoft Teams. The platform only schedules.'),
     ('Is my data secure?', 'Yes. Passwords are SHA-256 hashed client-side. No patient data or IP documents are stored. The platform is GDPR-compliant.'),
     ('How do I become an Admin?', 'Register with admin@healthai.edu to be auto-promoted. Subsequent admins must be granted by an existing admin.'),
     ('I see a recovery screen instead of the page', 'A render error was caught by the GlobalErrorBoundary. Click the recovery button to reload; if it persists, contact support with the URL.'),
     ('My session keeps expiring', 'A 30-minute inactivity timeout signs you out automatically. Sign in again to continue.')])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 14. CONTACT & SUPPORT
# ══════════════════════════════════════════════════════════════════
doc.add_heading('14. Contact & Support', level=1)
doc.add_paragraph(
    'For platform support, account issues, or to request administrator action please contact your '
    'institution\'s designated HEALTH AI administrator. For project-level questions about this '
    'release, the development team can be reached at the email addresses below.'
)
tbl(['Member', 'Role', 'Contact'],
    [('Cem Özal', 'Full-stack development', 'team@healthai.edu (placeholder)'),
     ('Emre Kurubaş', 'Frontend & UI/UX', 'team@healthai.edu (placeholder)'),
     ('Hasabu Can Eltayeb', 'Backend & Firebase', 'team@healthai.edu (placeholder)'),
     ('Sertaç Ataç', 'QA & documentation', 'team@healthai.edu (placeholder)')])

doc.add_paragraph(
    'Course / project ownership: SENG 384 — Software Engineering. Group: Health Shield. Document '
    'Version 3.1 — May 1, 2026.'
)

# ═══════════════ SAVE ═══════════════
output = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'HEALTH_AI_User_Guide.docx')
doc.save(output)
print(f'User Guide saved to: {output}')
